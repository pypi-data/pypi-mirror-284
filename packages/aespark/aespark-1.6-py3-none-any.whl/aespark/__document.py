from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.shared import Cm
import unicodedata
import pkg_resources
import pandas as pd
from docx import Document
from .__dataclean import *
import warnings
warnings.filterwarnings('ignore')

dist = pkg_resources.get_distribution("aespark")

class MyDocument():
    '''
    功能简介：
        读取模板文档创建空白文档供操作；
    参数解释：
        url :   可选 word模板路径；
    '''
    clolist = []  # 向doc增加表格时，格式化列名在此列表内的列内容，将金额规范为标准格式输出
    def __init__(self, url: str = '', drop: bool = True):
        if url == '':
            url = f'{dist.location}/aespark/static/word.docx'
        self.__doc = Document(url)
        if drop:
            self.new()

    def add_text(self, string: str = '未指定插入内容', level: int = -1):
        '''
        功能简介：
            在docx文档末尾增加新的内容，可以是标题或者段落；
        参数解释：
            string    :     新增的内容；
            level     :     内容格式，默认为段落，其余正整数为标题、对于标题等级；
        '''
        if level == -1:
            self.__doc.add_paragraph(string)
        else:
            self.__doc.add_heading(string, level=level)

    def add_pic(self, url:str, scale:float=1):
        '''
        功能简介：
            在docx文档末尾插入图片；
        参数解释：
            url 图片文件路径；
            scale   图片缩放比例，默认不缩放；例如 3：放大3倍，0.5：缩小50%；
        '''
        picture = self.__doc.add_picture(url)
        picture.width = int(picture.width * scale)
        picture.height = int(picture.height * scale)

    def add_table(self, df: pd.DataFrame, fontsize: int = 10):
        '''
        功能简介：
            在docx文档末尾插入新表格；
        参数解释：
            df          :   需插入的表格；
            fontsize    :   可选 表格内字体大小；
        '''
        def get_fontwidth(text):
            count = 0
            for c in text:
                if unicodedata.east_asian_width(c) in 'FWA':
                    count += 2
                else:
                    count += 0.83
            return count

        def set_tablewidth(table, widths):
            """表格分别设置列宽，单位为Cm"""
            for x, width in enumerate(widths):
                for cell in table.columns[x].cells:
                    cell.width = Cm(width)

        cloli = df.columns.to_list()
        if len(set(cloli)) != len(cloli):
            newcloli = []
            for i in cloli:
                if i not in newcloli:
                    newcloli.append(i)
                else:
                    for j in range(1,100):
                        if i+str(j) not in newcloli:
                            newcloli.append(i+str(j))
                            break
            df.columns = newcloli

        for col in self.clolist:
            if col in df.columns.to_list():
                df[col] = df[col].apply(lambda x: money_write(x, all=True))

        table = self.__doc.add_table(df.shape[0]+1, df.shape[1], style='Table Grid')
        table.style.font.size = Pt(fontsize)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i+1, j).text = str(df.values[i, j])
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        widlist = []
        for i in df.columns:
            wid = get_fontwidth(str(i))
            for index in df.index:
                lins = get_fontwidth(str(df[i][index]))
                wid = lins if lins > wid else wid
            widlist.append(wid)
        widlist = [i/sum(widlist)*15.26 for i in widlist]
        set_tablewidth(table, widlist)  # 共15.26

    def add_cloname(self, tt:str|list):
        '''
        功能介绍：
            添加需要进行金额格式化的列名；
        '''
        if type(tt) is str:
            self.clolist.append(tt)
        elif  type(tt) is list:
            self.clolist += tt
        self.clolist = list(set(self.clolist))

    def save(self, url: str = f"./未设置保存名称.docx"):
        '''
        功能简介：
            保存docx文档；
        参数解释：
            url :   str.    保存路径（含文件名及后缀），默认当前程序位置；
        '''
        self.__doc.save(url)

    def move_pars(self):
        '''
        功能简介：
            移除所有段落；
        '''
        for par in self.__doc.paragraphs:
            par._element.getparent().remove(par._element)

    def move_tables(self):
        '''
        功能简介：
            移除所有表格；
        '''
        for tab in self.__doc.tables:
            tab._element.getparent().remove(tab._element)

    def new(self):
        '''
        功能简介：
            清空文档；
        '''
        self.move_pars()
        self.move_tables()