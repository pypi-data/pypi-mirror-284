import re
from docx import Document
from pathlib import Path
from tqdm import tqdm
import pandas as pd
import time
import numpy as np
import warnings
warnings.filterwarnings('ignore')

def dc_dochide(docxurl:str, nameurl:str='', many:bool=False):
    '''
    功能简介：
        word文档脱敏；
    参数解释：
        docxurl 需要脱敏的文档路径
        nameurl 需要脱敏的名字txt文件路径，不传入则仅对数字脱敏
        many 是否要一次性操作多个docx文档，默认false，如果设置为true，docxurl需要传入文件夹路径
    文档解释：
        需要一个记录需要脱敏名称的txt文件，多个名字用空格隔开或者换行的方式写入；
    '''
    if nameurl != '':
        nlist = []
        name = open(nameurl,'r').readlines()
        for i in name:
            i = i.replace('\n','')
            nlist.append(i.split())
        nlist = [item for sublist in nlist for item in sublist]
        nlist = list(set(nlist))

    num_regex = re.compile(r'\d{10,}')

    def tuomin(url):
        doc = Document(url)
        for i in range(len(doc.paragraphs)):
            text = doc.paragraphs[i].text
            if nameurl != '':
                for x in nlist:
                    new_name = x[0] + '某' + x[-1]
                    text = text.replace(x, new_name)
            for num_match in num_regex.finditer(text):
                start = num_match.start()
                end = num_match.end()
                text = text[:start+3] + '8'*(end-start-7) + text[end-4:]
            doc.paragraphs[i].text = text
        for i in range(len(doc.tables)):
            for j in range(len(doc.tables[i].rows)):
                for x in range(len(doc.tables[i].rows[j].cells)):
                    text = doc.tables[i].rows[j].cells[x].text
                    if nameurl != '':
                        for l in nlist:
                            new_name = l[0] + '某' + l[-1]
                            text = text.replace(l, new_name)
                    for num_match in num_regex.finditer(text):
                        start = num_match.start()
                        end = num_match.end()
                        text = text[:start+3] + '8'*(end-start-7) + text[end-4:]
                    doc.tables[i].rows[j].cells[x].text = text

        doc.save(url[:url.rindex('.')]+'-脱敏完成'+url[url.rindex('.'):])

    if many:
        docxs = [str(i) for i in Path(docxurl).rglob('*.docx')]  # 获取目标文件夹下的所有文件路径
        for i in tqdm(docxs, desc='文档脱敏'):
            tuomin(i)
    else:
        tuomin(docxurl)

def money_write(num:float|int|str, len:int=2, all:bool=False):
    '''
    输出格式化金额，不成功则返回原始内容
    num :   数值
    len :   保留小数位数，默认2
    all :   是否要全写（全写：2345亿6123万4189.31， 简写：2345.61亿）
    '''
    try: 
        num = float(num)
        if all:
            if num<10000: return (("{:."+str(len)+"f}").format(num))
            if num<100000000: return f'{int(num/10000)}万{(("{:."+str(len)+"f}").format(num%10000))}'
            else: return f'{int(num/100000000)}亿{int(num%100000000/10000)}万{(("{:."+str(len)+"f}").format(num%10000))}'
        else:
            if num<10000: return (("{:."+str(len)+"f}").format(num)).strip('0.')
            if num<10000000: return (("{:."+str(len)+"f}").format(num/10000)).strip('0.')+'万'
            if num<100000000: return (("{:."+str(len)+"f}").format(num/10000000)).strip('0.')+'千万'
            else: return (("{:."+str(len)+"f}").format(num/100000000)).strip('0.')+'亿'
    except: return num

def dc_addtt(df: pd.DataFrame):
    '''
        为表格中超过15位的纯数字内容添加\t防止精度丢失
    '''
    df.columns = [str(i)+'\t' if str(i).isdigit() and len(str(i)) > 15 else i for i in df.columns]
    df = df.applymap(lambda x: str(x)+'\t' if str(x).isdigit() and len(str(x)) > 15 else x)

def dc_delerrchar(chars: str | pd.DataFrame, df: bool = False):
    '''
        清除不可见字符；
        chars 可传字符可传表，默认传的字符
        df 如果要传dataframe，该项参数需要填写为 True
    '''
    if df: chars = chars.applymap(lambda x: re.sub(u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', x) if type(x) == str else x)
    else:
        if type(chars) == str: chars = re.sub(u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', chars)
        return chars

def dc_inandout(str: str):
    '''
        统一借贷标志；
        需要清洗的字符，建议配合pandas.apply使用；
    当前可清洗内容：
        出 = ['借', '出', '支出', '付', 'D']；
        进 = ['贷', '进', '收入', '收', 'C']；
    如果发现了新的借贷标志可以进行添加；
    '''
    jie = ['借', '出', '支出', '付', 'D']
    dai = ['贷', '进', '收入', '收', 'C']
    return '出' if str in jie else '进' if str in dai else '其他'

def __try(timestr: str, format: str):
    '''
    功能简介：
        格式化时间格式；
    所需参数：
        timestr 需要格式化的字符串；
        format 字符串的格式（%Y年 %m月 %d日 %H时 %M分 %S秒）；
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）；
        若清洗失败则会返回False；
    '''
    timestr = str(timestr)
    try:
        timeStruct = time.strptime(timestr, format)
        times = time.strftime("%Y.%m.%d %H:%M:%S", timeStruct)
        return times
    except: return False

def dc_time(timestr: str):
    '''
    功能简介：
        兼容格式，批量格式化时间格式；
    所需参数：
        timestr 需要格式化的字符串（建议配合pandas.apply使用）；
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）；
        若清洗失败则会返回 nan；
    '''
    timestr = str(timestr)
    if timestr.isdigit():
        if len(timestr) == 14:
            times = __try(timestr, '%Y%m%d%H%M%S')
        elif len(timestr) == 12:
            times = __try(timestr, '%Y%m%d%H%M')
        elif len(timestr) == 8:
            times = __try(timestr, '%Y%m%d')
        else:
            times = __try(timestr, '%Y%m%d%H%M%S')
            if times is False:
                times = __try(timestr, '%Y%m%d%H%M')
            if times is False:
                times = __try(timestr, '%Y%m%d')
    else:
        if '-' in timestr:
            s = '-'
        elif '/' in timestr:
            s = '/'
        elif '.' in timestr:
            s = '.'
        else:
            s = ''
        times = __try(timestr, f'%Y{s}%m{s}%d %H:%M:%S')
        if times is False:
            times = __try(timestr, f'%Y{s}%m{s}%d %H:%M')

        if times is False:
            times = __try(timestr, f'%Y{s}%m{s}%d')

        if times is False and len(timestr) == 26:  # 2016-01-21-21.17.03.704713
            times = __try(timestr[:-7], '%Y-%m-%d-%H.%M.%S')

    return times if times else np.nan