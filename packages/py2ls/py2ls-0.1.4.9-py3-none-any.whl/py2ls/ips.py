from scipy.ndimage import convolve1d
from scipy.signal import savgol_filter
import pingouin as pg
from scipy import stats

import numpy as np
import pandas as pd

import json
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.ticker as tck 
from mpl_toolkits.mplot3d import Axes3D
import seaborn as sns

import sys, os,shutil,re, yaml,json
from cycler import cycler
import time
from dateutil import parser
from datetime import datetime

from PIL import Image,ImageEnhance, ImageOps,ImageFilter
from rembg import remove,new_session

import docx
from fpdf import FPDF 
from lxml import etree 
from docx import Document
from PyPDF2 import PdfReader
from pdf2image import convert_from_path, pdfinfo_from_path
from nltk.tokenize import sent_tokenize, word_tokenize
import nltk  # nltk.download("punkt")
from docx2pdf import convert
import img2pdf as image2pdf
import nbformat
from nbconvert import MarkdownExporter

from itertools import pairwise
from box import Box, BoxList
from numerizer import numerize
from tqdm import tqdm
import mimetypes
from pprint import pp
from collections import Counter
from fuzzywuzzy import fuzz,process
from langdetect import detect
from duckduckgo_search import DDGS

from py2ls import netfinder

dir_save='/Users/macjianfeng/Dropbox/Downloads/'

def rm_folder(folder_path, verbose=True):
    try:
        shutil.rmtree(folder_path)
        if verbose:
            print(f'Successfully deleted {folder_path}')
    except Exception as e:
        if verbose:
            print(f'Failed to delete {folder_path}. Reason: {e}')

def fremove(path, verbose=True):
    """
    Remove a folder and all its contents or a single file.
    Parameters:
    path (str): The path to the folder or file to remove.
    verbose (bool): If True, print success or failure messages. Default is True.
    """
    try:
        if os.path.isdir(path):
            shutil.rmtree(path)
            if verbose:
                print(f'Successfully deleted folder {path}')
        elif os.path.isfile(path):
            os.remove(path)
            if verbose:
                print(f'Successfully deleted file {path}')
        else:
            if verbose:
                print(f'Path {path} does not exist')
    except Exception as e:
        if verbose:
            print(f'Failed to delete {path}. Reason: {e}')


def get_cwd(verbose:bool = True):
    """
    get_cwd: to get the current working directory
    Args:
        verbose (bool, optional): to show which function is use. Defaults to True.
    """
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        if verbose:
            print("os.path.dirname(os.path.abspath(__file__)):", script_dir)
    except NameError:
        # This works in an interactive environment (like a Jupyter notebook)
        script_dir = os.getcwd()
        if verbose:
            print("os.getcwd():", script_dir) 

def search(query, limit=5, kind='text', output='df',verbose=False,download=True, dir_save=dir_save):
    from duckduckgo_search import DDGS
    if 'te' in kind.lower():
        results = DDGS().text(query, max_results=limit)
        res=pd.DataFrame(results)
        res.rename(columns={"href":"links"},inplace=True)
    if verbose:
        print(f'searching "{query}": got the results below\n{res}')
    if download:
        try:
            netfinder.downloader(url=res.links.tolist(), dir_save=dir_save, verbose=verbose)
        except:
            if verbose:
                print(f"failed link")
    return res

def echo(*args,**kwargs):
    """
    query, model="gpt", verbose=True, log=True, dir_save=dir_save
    a ai chat tool
    Args:
        query (str): _description_
        model (str, optional): _description_. Defaults to "gpt".
        verbose (bool, optional): _description_. Defaults to True.
        log (bool, optional): _description_. Defaults to True.
        dir_save (str, path, optional): _description_. Defaults to dir_save.

    Returns:
        str: the answer from ai
    """
    global dir_save
    
    query=None
    model=kwargs.get('model', 'gpt')
    verbose=kwargs.get('verbose', True)
    log=kwargs.get('log', True)
    dir_save=kwargs.get('dir_save', dir_save)
    for arg in args:
        if isinstance(arg, str):
            if os.path.isdir(arg):
                dir_save = arg
            # elif os.path.isfile(arg):
            #     dir_save = dirname(arg)
            elif len(arg) <= 5:
                model = arg
            else:
                query = arg
        elif isinstance(arg, dict):
            verbose = arg.get("verbose", verbose)
            log = arg.get("log", log)
    def is_in_any(str_candi_short, str_full, ignore_case=True):
        if isinstance(str_candi_short, str):
            str_candi_short=[str_candi_short]
        res_bool=[]
        if ignore_case:
            [res_bool.append(i in str_full.lower())  for i in str_candi_short ]
        else:
            [res_bool.append(i in str_full)  for i in str_candi_short ]
        return any(res_bool)
    def valid_mod_name(str_fly):
        if is_in_any(str_fly, "claude-3-haiku"):
            return "claude-3-haiku"
        elif is_in_any(str_fly, "gpt-3.5"):
            return "gpt-3.5"
        elif is_in_any(str_fly, "llama-3-70b"):
            return "llama-3-70b"
        elif is_in_any(str_fly, "mixtral-8x7b"):
            return "mixtral-8x7b"
        else:
            print(f"not support your model{model}, supported models: 'claude','gpt(default)', 'llama','mixtral'")
            return "gpt-3.5" # default model
    model_valid = valid_mod_name(model)
    res=DDGS().chat(query, model=model_valid)
    if verbose:
        pp(res)
    if log:
        dt_str=datetime.fromtimestamp(time.time()).strftime('%Y-%m-%d_%H:%M:%S')
        res_ = f"\n\n####Q:{query}\n\n#####Ans:{dt_str}\n\n>{res}\n"
        if bool(os.path.basename(dir_save)):
            fpath = dir_save
        else:
            os.makedirs(dir_save, exist_ok=True)
            fpath = os.path.join(dir_save, f"log_ai.md")
        fupdate(fpath=fpath,content=res_)
        print(f"log file:{fpath}")
    return res

def chat(*args, **kwargs):
    if len(args) == 1 and isinstance(args[0], str):
        kwargs['query'] = args[0]
    return echo(**kwargs)

def ai(*args, **kwargs):
    if len(args) == 1 and isinstance(args[0], str):
        kwargs['query'] = args[0]
    return echo(**kwargs)

def detect_lang(text, output='lang',verbose=True):
    dir_curr_script=os.path.dirname(os.path.abspath(__file__))
    dir_lang_code=dir_curr_script+"/data/lang_code_iso639.json"
    print(dir_curr_script,os.getcwd(),dir_lang_code)
    lang_code_iso639=fload(dir_lang_code)
    l_lang,l_code = [],[]
    [[l_lang.append(v),l_code.append(k)] for v,k in lang_code_iso639.items()]
    try:
        if is_text(text):
            code_detect=detect(text)
            if 'c' in output.lower(): # return code
                return l_code[strcmp(code_detect,l_code, verbose=verbose)[1]]
            else:
                return l_lang[strcmp(code_detect,l_code, verbose=verbose)[1]]
        else:
            print(f"{text} is not supported")
            return 'no'
    except:
        return 'no'

def is_text(s):
    has_alpha = any(char.isalpha() for char in s)
    has_non_alpha = any(not char.isalpha() for char in s)
    # no_special = not re.search(r'[^A-Za-z0-9\s]', s)
    return has_alpha and has_non_alpha

def strcmp(search_term, candidates, ignore_case=True, verbose=True, scorer='WR'):
    """
    Compares a search term with a list of candidate strings and finds the best match based on similarity score.

    Parameters:
    search_term (str): The term to be searched for.
    candidates (list of str): A list of candidate strings to compare against the search term.
    ignore_case (bool): If True, the comparison ignores case differences.
    verbose (bool): If True, prints the similarity score and the best match.

    Returns:
    tuple: A tuple containing the best match and its index in the candidates list.
    """
    def to_lower(s, ignore_case=True):
        #Converts a string or list of strings to lowercase if ignore_case is True.
        if ignore_case:
            if isinstance(s, str):
                return s.lower()
            elif isinstance(s, list):
                return [elem.lower() for elem in s]
        return s
    str1_,str2_ = to_lower(search_term, ignore_case),to_lower(candidates, ignore_case)
    if isinstance(str2_, list):
        if 'part' in scorer.lower():
            similarity_scores = [fuzz.partial_ratio(str1_, word) for word in str2_]
        elif 'W' in scorer.lower():
            similarity_scores = [fuzz.WRatio(str1_, word) for word in str2_]
        elif 'Ratio' in scorer.lower():
            similarity_scores = [fuzz.Ratio(str1_, word) for word in str2_]
        else:
            similarity_scores = [fuzz.WRatio(str1_, word) for word in str2_]
        best_match_index = similarity_scores.index(max(similarity_scores))
        best_match_score = similarity_scores[best_match_index]
    else:
        best_match_index = 0
        if 'part' in scorer.lower():
            best_match_score = fuzz.partial_ratio(str1_, str2_)
        elif 'W' in scorer.lower():
            best_match_score = fuzz.WRatio(str1_, str2_)
        elif 'Ratio' in scorer.lower():
            best_match_score = fuzz.Ratio(str1_, str2_)
        else:
            best_match_score = fuzz.WRatio(str1_, str2_)
    if verbose:
        print(f"\nbest_match is: {candidates[best_match_index],best_match_score}")
        best_match = process.extract(search_term, candidates)
        print(f"建议: {best_match}")
    return candidates[best_match_index], best_match_index

# Example usaged
# str1 = "plos biology"
# str2 = ['PLoS Computational Biology', 'PLOS BIOLOGY']
# best_match, idx = strcmp(str1, str2, ignore_case=1)

def counter(list_, verbose=True):
    c = Counter(list_)
    # Print the name counts
    for item, count in c.items():
        if verbose:
            print(f"{item}: {count}")
    return c
# usage:
# print(f"Return an iterator over elements repeating each as many times as its count:\n{sorted(c.elements())}")
# print(f"Return a list of the n most common elements:\n{c.most_common()}")
# print(f"Compute the sum of the counts:\n{c.total()}")



def str2time(time_str, fmt='24'):
    """
    Convert a time string into the specified format.
    Parameters:
    - time_str (str): The time string to be converted.
    - fmt (str): The format to convert the time to. Defaults to '%H:%M:%S'.
    Returns:
        %I represents the hour in 12-hour format.
        %H represents the hour in 24-hour format (00 through 23).
        %M represents the minute.
        %S represents the second.
        %p represents AM or PM.
    - str: The converted time string.
    """
    def time_len_corr(time_str):
        time_str_= ssplit(time_str,by=[':'," ","digital_num"]) if ':' in time_str else None
        time_str_split=[]
        [time_str_split.append(i) for i in time_str_ if is_num(i)] 
        if time_str_split:
            if len(time_str_split)==2:
                H,M=time_str_split
                time_str_full=H+":"+M+":00"
            elif len(time_str_split)==3:
                H,M,S=time_str_split
                time_str_full=H+":"+M+":"+S
        else:
            time_str_full=time_str_
        if 'am' in time_str.lower():
            time_str_full+=" AM" 
        elif "pm"in time_str.lower():
            time_str_full +=" PM"  
        return time_str_full
    if '12' in fmt:
        fmt = "%I:%M:%S %p"
    elif '24' in fmt:
        fmt = "%H:%M:%S"

    try:
        # Try to parse the time string assuming it could be in 24-hour or 12-hour format
        time_obj = datetime.strptime(time_len_corr(time_str), '%H:%M:%S')
    except ValueError:
        try:
            time_obj = datetime.strptime(time_len_corr(time_str), '%I:%M:%S %p')
        except ValueError as e:
            raise ValueError(f"Unable to parse time string: {time_str}. Error: {e}")
    
    # Format the time object to the desired output format
    formatted_time = time_obj.strftime(fmt)
    return formatted_time

# # Example usage:
# time_str1 = "14:30:45"
# time_str2 = "02:30:45 PM"

# formatted_time1 = str2time(time_str1, fmt='12')  # Convert to 12-hour format
# formatted_time2 = str2time(time_str2, fmt='24')    # Convert to 24-hour format

# print(formatted_time1)  # Output: 02:30:45 PM
# print(formatted_time2)  # Output: 14:30:45

def str2date(date_str, fmt='%Y-%m-%d_%H:%M:%S'):
    """
    Convert a date string into the specified format.
    Parameters:
    - date_str (str): The date string to be converted.
    - fmt (str): The format to convert the date to. Defaults to '%Y%m%d'.
    Returns:
    - str: The converted date string.
    """
    try:
        date_obj = parser.parse(date_str)
    except ValueError as e:
        raise ValueError(f"Unable to parse date string: {date_str}. Error: {e}")
    # Format the date object to the desired output format
    formatted_date = date_obj.strftime(fmt)
    return formatted_date
# str1=str2date(num2str(20240625),fmt="%a %d-%B-%Y")
# print(str1)
# str2=str2num(str2date(str1,fmt='%a %Y%m%d'))
# print(str2)

def str2num(s, *args):
    delimiter = None
    round_digits = None
    for arg in args:
        if isinstance(arg, str):
            delimiter = arg
        elif isinstance(arg, int):
            round_digits = arg
    try:
        num = int(s)
    except ValueError:
        try:
            num = float(s)
        except ValueError:
            try:
                numerized = numerize(s)
                num = int(numerized) if '.' not in numerized else float(numerized)
            except Exception as e:
                # Attempt to handle multiple number segments
                try:
                    number_segments = ssplit(s,by='number_strings')
                    nums = []
                    for segment in number_segments:
                        try:
                            nums.append(str2num(segment))
                        except ValueError:
                            continue
                    if len(nums) == 1:
                        num = nums[0]
                    else:
                        raise ValueError("Multiple number segments found, cannot determine single numeric value")
                except Exception as e:
                    raise ValueError(f"Cannot convert {s} to a number: {e}")

    # Apply rounding if specified
    if round_digits is not None:
        num_adj = num + 0.00000000001  # Ensure precise rounding
        num = round(num_adj, round_digits)

    # Apply delimiter formatting if specified
    if delimiter is not None:
        num_str = f"{num:,}".replace(",", delimiter)
        return num_str

    return num
# Examples
# print(str2num("123"))                # Output: 123
# print(str2num("123.456", 2))         # Output: 123.46
# print(str2num("one hundred and twenty three"))  # Output: 123
# print(str2num("seven million"))      # Output: 7000000
# print(str2num('one thousand thirty one',','))  # Output: 1,031
# print(str2num("12345.6789", ","))    # Output: 12,345.6789
# print(str2num("12345.6789", " ", 2)) # Output: 12 345.68
# print(str2num('111113.34555',3,',')) # Output: 111,113.346
# print(str2num("123.55555 sec miniuets",3)) # Output: 1.3
def num2str(num, *args):
    delimiter = None
    round_digits = None

    # Parse additional arguments
    for arg in args:
        if isinstance(arg, str):
            delimiter = arg
        elif isinstance(arg, int):
            round_digits = arg

    # Apply rounding if specified
    if round_digits is not None:
        num = round(num, round_digits)

    # Convert number to string
    num_str = f"{num}"

    # Apply delimiter if specified
    if delimiter is not None:
        num_str = num_str.replace(".", ",")  # Replace decimal point with comma
        num_str_parts = num_str.split(",")
        if len(num_str_parts) > 1:
            integer_part = num_str_parts[0]
            decimal_part = num_str_parts[1]
            integer_part = "{:,}".format(int(integer_part))
            num_str = integer_part + "." + decimal_part
        else:
            num_str = "{:,}".format(int(num_str_parts[0]))

    return num_str
# Examples
# print(num2str(123),type(num2str(123)))                # Output: "123"
# print(num2str(123.456, 2),type(num2str(123.456, 2)))         # Output: "123.46"
# print(num2str(7000.125, 2),type(num2str(7000.125, 2)))        # Output: "7000.13"
# print(num2str(12345.6789, ","),type(num2str(12345.6789, ",")))    # Output: "12,345.6789"
# print(num2str(7000.00, ","),type(num2str(7000.00, ",")))       # Output: "7,000.00"
def sreplace(*args,**kwargs):
    """
    sreplace(text, by=None, robust=True)
    Replace specified substrings in the input text with provided replacements.
    Args:
        text (str): The input text where replacements will be made.
        by (dict, optional): A dictionary containing substrings to be replaced as keys
            and their corresponding replacements as values. Defaults to {".com": "..come", "\n": " ", "\t": " ", "  ": " "}.
        robust (bool, optional): If True, additional default replacements for newline and tab characters will be applied.
                                Default is False.
    Returns:
        str: The text after replacements have been made.
    """
    text = None
    by = kwargs.get('by', None)
    robust = kwargs.get('robust', True)
    
    for arg in args:
        if isinstance(arg,str):
            text=arg
        elif isinstance(arg,dict):
            by=arg
        elif isinstance(arg,bool):
            robust=arg
        else:
            Error(f"{type(arg)} is not supported")
    
    # Default replacements for newline and tab characters
    default_replacements = {
        "\a": "",
        "\b": "",
        "\f": "",
        "\n": "",
        "\r": "",
        "\t": "",
        "\v": "",
        "\\": "",  # Corrected here
        # "\?": "",
        "�": "",
        "\\x": "",  # Corrected here
        "\\x hhhh": "",
        "\\ ooo": "",  # Corrected here
        "\xa0": "",
        "  ": " ",
    }

    # If dict_replace is None, use the default dictionary
    if by is None:
        by = {}
    # If robust is True, update the dictionary with default replacements
    if robust:
        by.update(default_replacements)

    # Iterate over each key-value pair in the dictionary and replace substrings accordingly
    for k, v in by.items():
        text = text.replace(k, v)
    return text
# usage:
# sreplace(text, by=dict(old_str='new_str'), robust=True)

def paper_size(paper_type_str='a4'):
    df=pd.DataFrame({'a0':[841,1189],'a1':[594,841],'a2':[420,594],'a3':[297,420],'a4':[210,297],'a5':[148,210],'a6':[105,148],'a7':[74,105],
                     'b0':[1028,1456],'b1':[707,1000],'b2':[514,728],'b3':[364,514],'b4':[257,364],'b5':[182,257],'b6':[128,182],
                     'letter': [215.9, 279.4],'legal':[215.9, 355.6],'business card':[85.6, 53.98],
                     'photo china passport':[33,48],'passport single':[125,88],'visa':[105,74],'sim':[25,15]})
    for name in df.columns:
        if paper_type_str in name.lower():
            paper_type=name
    if not paper_type:
        paper_type='a4' # default
    return df[paper_type].tolist()

def docx2pdf(dir_docx, dir_pdf=None):
    if dir_pdf:
        convert(dir_docx,dir_pdf)
    else:
        convert(dir_docx)

def img2pdf(dir_img, kind="jpeg",page=None, dir_save=None, page_size="a4", dpi=300):
    def mm_to_point(size):
        return (image2pdf.mm_to_pt(size[0]),image2pdf.mm_to_pt(size[1]))
    def set_dpi(x):
        dpix=dpiy=x
        return image2pdf.get_fixed_dpi_layout_fun((dpix, dpiy))
    if not kind.startswith("."):
        kind="."+kind
    if dir_save is None:
        dir_save = dir_img.replace(kind,'.pdf')
    imgs = [] 
    if os.path.isdir(dir_img):
        if not dir_save.endswith(".pdf"):
            dir_save+="#merged_img2pdf.pdf"
        if page is None:
            select_range = listdir(dir_img,kind=kind).fpath
        else:
            if not isinstance(page, (np.ndarray,list,range)):
                page=[page]
            select_range = listdir(dir_img,kind=kind)['fpath'][page]
        for fname in select_range:
            if not fname.endswith(kind):
                continue
            path = os.path.join(dir_img, fname)
            if os.path.isdir(path):
                continue
            imgs.append(path)
    else:
        imgs=[os.path.isdir(dir_img),dir_img]

    if page_size:
        if isinstance(page_size,str):
            pdf_in_mm=mm_to_point(paper_size(page_size))
        else:
            print("default: page_size = (210,297)")
            pdf_in_mm=mm_to_point(page_size)
            print(f"page size was set to {page_size}")
        p_size= image2pdf.get_layout_fun(pdf_in_mm)
    else:
        p_size = set_dpi(dpi)
    with open(dir_save,"wb") as f:
        f.write(image2pdf.convert(imgs, layout_fun=p_size))
# usage:
# dir_img="/Users/macjianfeng/Dropbox/00-Personal/2015-History/2012-2015_兰州大学/120901-大学课件/生物统计学 陆卫/复习题/"
# img2pdf(dir_img,kind='tif', page=range(3,7,2))
def ssplit(text, by="space", verbose=False, **kws):
    if isinstance(text, list):
        nested_list= [ssplit(i,by=by,verbose=verbose,**kws) for i in text]
        flat_list = [item for sublist in nested_list for item in sublist]
        return flat_list
    def split_by_word_length(text, length):
        return [word for word in text.split() if len(word) == length]

    def split_by_multiple_delimiters(text, delimiters):
        regex_pattern = "|".join(map(re.escape, delimiters))
        return re.split(regex_pattern, text)

    def split_by_camel_case(text):
        return re.findall(r"[A-Z](?:[a-z]+|[A-Z]*(?=[A-Z]|$))", text)

    def split_at_upper_fl_lower(text):
        return re.findall(r"[A-Z](?:[a-z]+|[A-Z]+(?=[A-Z]|$))", text)

    def split_at_lower_fl_upper(text):
        split_text = re.split(r"(?<=[a-z])(?=[A-Z])", text)
        return split_text

    def split_at_upper(text):
        split_text = re.split(r"(?=[A-Z])", text)
        split_text = [part for part in split_text if part]
        return split_text

    def split_by_regex_lookahead(text, pattern):
        return re.split(f'(?<={pattern})', text)
    
    def split_by_regex_end(text, pattern):
        return re.split(f'(?={pattern})', text)

    # def split_by_sentence_endings(text):
    #     return re.split(r"(?<=[.!?])", text)
    def split_non_ascii(text):
        # return re.split(r"([^\x00-\x7F\w\s,.!?:\"'()\-]+)", text)
        # return re.split(r"[^\x00-\x7F]+", text)
        return re.split(r"([^\x00-\x7F]+)", text)
    def split_by_consecutive_non_alphanumeric(text):
        return re.split(r"\W+", text)

    def split_by_fixed_length_chunks(text, length):
        return [text[i : i + length] for i in range(0, len(text), length)]
    def split_by_sent_num(text,n=10):
        # split text into sentences
        text_split_by_sent=sent_tokenize(text)
        cut_loc_array=np.arange(0,len(text_split_by_sent),n)
        if cut_loc_array[-1]!=len(text_split_by_sent):
            cut_loc=np.append(cut_loc_array,len(text_split_by_sent))
        else:
            cut_loc = cut_loc_array
        # get text in section (e.g., every 10 sentences)
        text_section=[]
        for i,j in pairwise(cut_loc):
            text_section.append(text_split_by_sent[i:j])
        return text_section
    def split_general(text, by, verbose=False, ignore_case=False):
        if ignore_case:
            if verbose:
                print(f"used {by} to split, ignore_case=True")
            pattern = re.compile(re.escape(by), re.IGNORECASE)
            split_text = pattern.split(text)
            return split_text
        else:
            if verbose:
                print(f"used {by} to split, ignore_case=False")
            return text.split(by)
    def reg_split(text, pattern):
        return re.split(pattern, text)
    if "sp" in by or "white" in by:
        if verbose:
            print(f"splited by space")
        return text.split()
    elif "word" in by and "len" in by:
        if verbose:
            print(f"split_by_word_length(text, length)")
        return split_by_word_length(text, **kws)  # split_by_word_length(text, length)
    # elif "," in by:
    #     if verbose:
    #         print(f"splited by ','")
    #     return text.split(",")
    elif isinstance(by, list):
        if verbose:
            print(f"split_by_multiple_delimiters: ['|','&']")
        return split_by_multiple_delimiters(text, by)
    elif all([("digi" in by or "num" in by),not 'sent' in by, not 'str' in by]):
        if verbose:
            print(f"splited by digital (numbers)")
        return re.split(r"(\d+)", text)
    elif all([("digi" in by or "num" in by), 'str' in by]):
        if verbose:
            print(f"Splitting by (number strings)")
        pattern = re.compile(r'\b((?:one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety|hundred|thousand|million|billion|trillion|and|[\d,]+(?:\.\d+)?)(?:[-\s]?(?:one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety|hundred|thousand|million|billion|trillion|and|[\d,]+(?:\.\d+)?))*)\b', re.IGNORECASE)
        return re.split(pattern, text)
    elif "pun" in by:
        if verbose:
            print(f"splited by 标点('.!?;')")
        return re.split(r"[.!?;]", text)
    elif "\n" in by or "li" in by:
        if verbose:
            print(f"splited by lines('\n')")
        return text.splitlines()
    elif "cam" in by:
        if verbose:
            print(f"splited by camel_case")
        return split_by_camel_case(text)
    elif "word" in by:
        if verbose:
            print(f"splited by word")
        return word_tokenize(text)
    elif "sen" in by and not 'num' in by:
        if verbose:
            print(f"splited by sentence")
        return sent_tokenize(text)
    elif 'sen' in by and 'num' in by:
        return split_by_sent_num(text,**kws)
    elif "cha" in by:
        if verbose:
            print(f"splited by chracters")
        return list(text)
    elif ("up" in by or "cap" in by) and "l" not in by:
        if verbose:
            print(f"splited by upper case")
        return split_at_upper(text)
    elif "u" in by and "l" in by:
        if by.find("u") < by.find("l"):
            if verbose:
                print(f"splited by upper followed by lower case")
            return split_at_upper_fl_lower(text)
        else:
            if verbose:
                print(f"splited by lower followed by upper case")
            return split_at_lower_fl_upper(text)
    elif "start" in by or "head" in by:
        if verbose:
            print(f"splited by lookahead")
        return split_by_regex_lookahead(text, **kws)
    elif "end" in by or "tail" in by:
        if verbose:
            print(f"splited by endings")
        return split_by_regex_end(text, **kws)
    elif "other" in by or "non_alp" in by:
        if verbose:
            print(f"splited by non_alphanumeric")
        return split_by_consecutive_non_alphanumeric(text)
    elif "len" in by:
        if verbose:
            print(f"splited by fixed length")
        return split_by_fixed_length_chunks(text, **kws)
    elif "re" in by or "cus" in by or "cos" in by:
        if verbose:
            print(f"splited by customed, re; => {by}")
        return reg_split(text, **kws)
    elif 'lang' in by or 'eng' in by:
        return split_non_ascii(text)
    else:
        return split_general(text, by, verbose=verbose, **kws)


def pdf2img(dir_pdf, dir_save=None, page=None, kind="png",verbose=True, **kws):
    df_dir_img_single_page = pd.DataFrame()
    dir_single_page = []
    if verbose:
        pp(pdfinfo_from_path(dir_pdf))
    if isinstance(page, tuple) and page:
        page = list(page)
    if isinstance(page,int):
        page=[page]
    if page is None:
        page = [pdfinfo_from_path(dir_pdf)["Pages"]]
    if len(page)==1 and page != pdfinfo_from_path(dir_pdf)["Pages"]:
        page=[page[0], page[0]]
    else:
        page=[1, page[0]]
    pages = convert_from_path(dir_pdf, first_page=page[0], last_page=page[1], **kws)
    if dir_save is None:
        dir_save = newfolder(dirname(dir_pdf), basename(dir_pdf).split(".")[0] + "_img")
    for i, page in enumerate(pages):
        if verbose: 
            print(f"processing page: {i+1}")
        if i < 9:
            dir_img_each_page = dir_save + f"page_0{i+1}.png"
        else:
            dir_img_each_page = dir_save + f"page_{i+1}.png"
        dir_single_page.append(dir_img_each_page)
        page.save(dir_img_each_page, kind.upper())
    df_dir_img_single_page["fpath"] = dir_single_page
    return df_dir_img_single_page

# dir_pdf = "/Users/macjianfeng/Dropbox/github/python/240308_Python Data Science Handbook.pdf"
# df_page = pdf2img(dir_pdf, page=[1, 5],dpi=300)



def fload(fpath, kind=None, **kwargs):
    """
    Load content from a file with specified file type.
    Parameters:
        fpath (str): The file path from which content will be loaded.
        kind (str): The file type to load. Supported options: 'docx', 'txt', 'md', 'html', 'json', 'yaml', 'xml', 'csv', 'xlsx', 'pdf'.
        **kwargs: Additional parameters for 'csv' and 'xlsx' file types.
    Returns:
        content: The content loaded from the file.
    """
    def load_txt_md(fpath):
        with open(fpath, "r") as file:
            content = file.read()
        return content

    def load_html(fpath):
        with open(fpath, "r") as file:
            content = file.read()
        return content

    def load_json(fpath):
        with open(fpath, "r") as file:
            content = json.load(file)
        return content

    def load_yaml(fpath):
        with open(fpath, "r") as file:
            content = yaml.safe_load(file)
        return content

    def load_xml(fpath):
        tree = etree.parse(fpath)
        root = tree.getroot()
        return etree.tostring(root, pretty_print=True).decode()

    def load_csv(fpath, **kwargs):
        df = pd.read_csv(fpath, **kwargs)
        return df

    def load_xlsx(fpath, **kwargs):
        df = pd.read_excel(fpath, **kwargs)
        return df
    def load_ipynb(fpath,**kwargs):
        as_version=kwargs.get("as_version",4)
        with open(fpath, "r") as file:
            nb = nbformat.read(file, as_version=as_version)
            md_exporter = MarkdownExporter()
            md_body, _ = md_exporter.from_notebook_node(nb)
        return md_body
    
    def load_pdf(fpath, page='all', verbose=False, **kwargs):
        """
            Parameters:
            fpath: The path to the PDF file to be loaded.
            page (optional): 
                Specifies which page or pages to extract text from. By default, it's set to "all", which means text from all 
                pages will be returned. It can also be an integer to specify a single page number or a list of integers to 
                specify multiple pages.
            verbose (optional): 
                If True, prints the total number of pages processed.
            Functionality:
            It initializes an empty dictionary text_dict to store page numbers as keys and their corresponding text as values.
            It iterates through each page of the PDF file using a for loop.
            For each page, it extracts the text using PyPDF2's extract_text() method and stores it in text_dict with the page number incremented by 1 as the key.
            If the page parameter is an integer, it converts it into a list containing that single page number to ensure consistency in handling.
            If the page parameter is a NumPy array, it converts it to a list using the tolist() method to ensure compatibility with list operations.
            If verbose is True, it prints the total number of pages processed.
            If page is a list, it combines the text of the specified pages into a single string combined_text and returns it.
            If page is set to "all", it returns the entire text_dict containing text of all pages.
            If page is an integer, it returns the text of the specified page number.
            If the specified page is not found, it returns the string "Page is not found".
        """
        text_dict = {}
        with open(fpath, "rb") as file:
            pdf_reader = PdfReader(file)
            num_pages = len(pdf_reader.pages)
            for page_num in range(num_pages):
                if verbose:
                    print(f"processing page {page_num}")
                page_ = pdf_reader.pages[page_num]
                text_dict[page_num + 1] = page_.extract_text()
        if isinstance(page, int):
            page = [page]
        elif isinstance(page, np.ndarray):
            page = page.tolist()
        if verbose:
            print(f"total pages: {page_num}")
        if isinstance(page, list):
            combined_text = ""
            for page_num in page:
                combined_text += text_dict.get(page_num, "")
            return combined_text
        elif "all" in page.lower():
            combined_text = ""
            for i in text_dict.values():
                combined_text += i
            return combined_text
        else:
            return text_dict.get(int(page), "Page is not found")

    def load_docx(fpath):
        doc = Document(fpath)
        content = [para.text for para in doc.paragraphs]
        return content 

    if kind is None:
        _, kind = os.path.splitext(fpath)
        kind = kind.lower()

    kind = kind.lstrip('.').lower()
    img_types=[ 'bmp','eps', 'gif', 'icns', 'ico', 'im', 'jpg','jpeg', 'jpeg2000','msp', 'pcx', 'png', 'ppm', 'sgi', 'spider', 'tga','tiff','webp',"json"]
    doc_types = ["docx", "txt", "md", "html", "json", "yaml", "xml", "csv", "xlsx", "pdf","ipynb"]
    supported_types = [*doc_types, *img_types]
    if kind not in supported_types:
        raise ValueError(f"Error:\n{kind} is not in the supported list {supported_types}")
    if kind == "docx":
        return load_docx(fpath)
    elif kind == "txt" or kind == "md":
        return load_txt_md(fpath)
    elif kind == "html":
        return load_html(fpath)
    elif kind == "json":
        return load_json(fpath)
    elif kind == "yaml":
        return load_yaml(fpath)
    elif kind == "xml":
        return load_xml(fpath)
    elif kind == "csv":
        return load_csv(fpath, **kwargs)
    elif kind == "xlsx":
        return load_xlsx(fpath, **kwargs)
    elif kind == "ipynb":
        return load_ipynb(fpath, **kwargs)
    elif kind == "pdf":
        print('usage:load_pdf(fpath, page="all", verbose=False)')
        return load_pdf(fpath, **kwargs)
    elif kind.lower() in img_types:
        print(f'Image ".{kind}" is loaded.')
        return load_img(fpath) 
    else:
        raise ValueError(f"Error:\n{kind} is not in the supported list {supported_types}")

# Example usage
# txt_content = fload('sample.txt')
# md_content = fload('sample.md')
# html_content = fload('sample.html')
# json_content = fload('sample.json')
# yaml_content = fload('sample.yaml')
# xml_content = fload('sample.xml')
# csv_content = fload('sample.csv')
# xlsx_content = fload('sample.xlsx')
# docx_content = fload('sample.docx')

def fupdate(fpath, content=None):
    """
    Update a file by adding new content at the top and moving the old content to the bottom.
    Parameters
    ----------
    fpath : str
        The file path where the content should be updated.
    content : str, optional
        The new content to add at the top of the file. If not provided, the function will not add any new content.
    Notes
    -----
    - If the file at `fpath` does not exist, it will be created.
    - The new content will be added at the top, followed by the old content of the file.
    """
    content = content or ""
    if os.path.exists(fpath):
        with open(fpath, 'r') as file:
            old_content = file.read()
    else:
        old_content = '' 
        
    with open(fpath, 'w') as file:
        file.write(content)
        file.write(old_content)
        
def fsave(
    fpath,
    content,
    kind=None,
    font_name="Times",
    font_size=10,
    spacing=6,
    mode='w',
    **kwargs,
):
    """
    Save content into a file with specified file type and formatting.
    Parameters:
        fpath (str): The file path where content will be saved.
        content (list of str or dict): The content to be saved, where each string represents a paragraph or a dictionary for tabular data.
        kind (str): The file type to save. Supported options: 'docx', 'txt', 'md', 'html', 'pdf', 'csv', 'xlsx', 'json', 'xml', 'yaml'.
        font_name (str): The font name for text formatting (only applicable for 'docx', 'html', and 'pdf').
        font_size (int): The font size for text formatting (only applicable for 'docx', 'html', and 'pdf').
        spacing (int): The space after each paragraph (only applicable for 'docx').
        **kwargs: Additional parameters for 'csv', 'xlsx', 'json', 'yaml' file types.
    Returns:
        None
    """
    def save_content(fpath, content, mode=mode): 
        with open(fpath, mode, encoding='utf-8') as file:
            file.write(content)


    def save_docx(fpath, content, font_name, font_size, spacing):
        if isinstance(content, str):
            content = content.split(". ")
        doc = docx.Document()
        for i, paragraph_text in enumerate(content):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(paragraph_text)
            font = run.font
            font.name = font_name
            font.size = docx.shared.Pt(font_size)
            if i != len(content) - 1:  # Add spacing for all but the last paragraph
                paragraph.space_after = docx.shared.Pt(spacing)
        doc.save(fpath)


    def save_txt_md(fpath, content, sep="\n",mode='w'):
            # Ensure content is a single string
        if isinstance(content, list):
            content = sep.join(content)
        save_content(fpath, sep.join(content),mode)


    def save_html(fpath, content, font_name, font_size,mode='w'):
        html_content = "<html><body>"
        for paragraph_text in content:
            html_content += f'<p style="font-family:{font_name}; font-size:{font_size}px;">{paragraph_text}</p>'
        html_content += "</body></html>"
        save_content(fpath, html_content,mode)


    def save_pdf(fpath, content, font_name, font_size):
        pdf = FPDF()
        pdf.add_page()
        # pdf.add_font('Arial','',r'/System/Library/Fonts/Supplemental/Arial.ttf',uni=True)
        pdf.set_font(font_name, '',font_size)
        for paragraph_text in content:
            pdf.multi_cell(0, 10, paragraph_text)
            pdf.ln(h = '')
        pdf.output(fpath,'F')


    def save_csv(fpath, data, **kwargs):
        df = pd.DataFrame(data)
        df.to_csv(fpath, **kwargs)


    def save_xlsx(fpath, data, **kwargs):
        df = pd.DataFrame(data)
        df.to_excel(fpath, **kwargs)

    def save_ipynb(fpath,data,**kwargs):
        # Split the content by code fences to distinguish between code and markdown
        parts = data.split('```')
        cells = []

        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Even index: markdown content
                cells.append(nbf.v4.new_markdown_cell(part.strip()))
            else:
                # Odd index: code content
                cells.append(nbf.v4.new_code_cell(part.strip()))
        # Create a new notebook
        nb = nbformat.v4.new_notebook()
        nb['cells'] = cells
        # Write the notebook to a file
        with open(fpath, 'w', encoding='utf-8') as ipynb_file:
            nbf.write(fpath, ipynb_file)
        
    # def save_json(fpath, data, **kwargs):
    #     with open(fpath, "w") as file:
    #         json.dump(data, file, **kwargs)

    def save_json(fpath_fname, var_dict_or_df): 
        with open(fpath_fname, "w") as f_json:
            # Check if var_dict_or_df is a DataFrame
            if isinstance(var_dict_or_df, pd.DataFrame):
                # Convert DataFrame to a list of dictionaries
                var_dict_or_df = var_dict_or_df.to_dict(orient="dict")
            
            # Check if var_dict_or_df is a dictionary
            if isinstance(var_dict_or_df, dict):
                # Convert NumPy arrays to lists
                for key, value in var_dict_or_df.items():
                    if isinstance(value, np.ndarray):
                        var_dict_or_df[key] = value.tolist()
            
            # Save the dictionary or list of dictionaries to a JSON file
            json.dump(var_dict_or_df, f_json, indent=4)
    # # Example usage:
    # sets = {"title": "mse_path_ MSE"}
    # jsonsave("/.json", sets)
    # # setss = jsonload("/.json")

    def save_yaml(fpath, data, **kwargs):
        with open(fpath, "w") as file:
            yaml.dump(data, file, **kwargs)


    def save_xml(fpath, data):
        root = etree.Element("root")
        if isinstance(data, dict):
            for key, val in data.items():
                child = etree.SubElement(root, key)
                child.text = str(val)
        else:
            raise ValueError("XML saving only supports dictionary data")
        tree = etree.ElementTree(root)
        tree.write(fpath, pretty_print=True, xml_declaration=True, encoding="UTF-8")

    if kind is None:
        _, kind = os.path.splitext(fpath)
        kind = kind.lower()

    kind = kind.lstrip(".").lower()

    if kind not in [
        "docx",
        "txt",
        "md",
        "html",
        "pdf",
        "csv",
        "xlsx",
        "json",
        "xml",
        "yaml",
        "ipynb"
    ]:
        print(
            f"Warning:\n{kind} is not in the supported list ['docx', 'txt', 'md', 'html', 'pdf', 'csv', 'xlsx', 'json', 'xml', 'yaml']"
        )

    if kind == "docx" or kind=="doc":
        save_docx(fpath, content, font_name, font_size, spacing)
    elif kind == "txt":
        save_txt_md(fpath, content, sep="",mode=mode)
    elif kind == "md":
        save_txt_md(fpath, content, sep="",mode=mode)
    elif kind == "html":
        save_html(fpath, content, font_name, font_size)
    elif kind == "pdf":
        save_pdf(fpath, content, font_name, font_size)
    elif kind == "csv":
        save_csv(fpath, content, **kwargs)
    elif kind == "xlsx":
        save_xlsx(fpath, content, **kwargs)
    elif kind == "json":
        save_json(fpath, content) 
    elif kind == "xml":
        save_xml(fpath, content) 
    elif kind == "yaml":
        save_yaml(fpath, content, **kwargs)
    elif kind == "ipynb":
        save_ipynb(fpath, content, **kwargs) 
    else:
        try: 
            netfinder.downloader(url=content, dir_save=dirname(fpath), kind=kind)
        except:
            print(
                f"Error:\n{kind} is not in the supported list ['docx', 'txt', 'md', 'html', 'pdf', 'csv', 'xlsx', 'json', 'xml', 'yaml']"
                )


# # Example usage
# text_content = ["Hello, this is a sample text file.", "This is the second paragraph."]
# tabular_content = {"Name": ["Alice", "Bob"], "Age": [24, 30]}
# json_content = {"name": "Alice", "age": 24}
# yaml_content = {"Name": "Alice", "Age": 24}
# xml_content = {"Name": "Alice", "Age": 24}
# dir_save = "/Users/macjianfeng/Dropbox/Downloads/"
# fsave(dir_save + "sample.txt", text_content)
# fsave(dir_save + "sample.md", text_content)
# fsave(dir_save + "sample.html", text_content)
# fsave(dir_save + "sample.pdf", text_content)
# fsave(dir_save + "sample.docx", text_content)
# fsave(dir_save + "sample.csv", tabular_content, index=False)
# fsave(dir_save + "sample.xlsx", tabular_content, sheet_name="Sheet1", index=False)
# fsave(dir_save + "sample.json", json_content, indent=4)
# fsave(dir_save + "sample.yaml", yaml_content)
# fsave(dir_save + "sample.xml", xml_content)

def addpath(fpath):
    sys.path.insert(0,dir)
def dirname(fpath):
    """
    dirname: Extracting Directory Name from a File Path
    Args:
        fpath (str): the file or directory path 
    Returns:
        str: directory, without filename
    """
    dirname_=os.path.dirname(fpath)
    if not dirname_.endswith('/'):
        dirname_=dirname_+"/"
    return dirname_

def dir_name(fpath): # same as "dirname"
    return dirname(fpath)
def basename(fpath):
    """
    basename: # Output: file.txt
    Args:
        fpath (str): the file or directory path 
    Returns:
        str: # Output: file.txt
    """
    return os.path.basename(fpath)
def flist(fpath, contains="all"):
    all_files = [os.path.join(fpath, f) for f in os.listdir(fpath) if os.path.isfile(os.path.join(fpath, f))]
    if isinstance(contains, list):
        filt_files = []
        for filter_ in contains:
            filt_files.extend(flist(fpath, filter_))
        return filt_files
    else:
        if 'all' in contains.lower():
            return all_files
        else:
            filt_files = [f for f in all_files if isa(f, contains)]
            return filt_files
def sort_kind(df, by="name", ascending=True):
    if df[by].dtype == 'object':  # Check if the column contains string values
        if ascending:
            sorted_index = df[by].str.lower().argsort()
        else:
            sorted_index = df[by].str.lower().argsort()[::-1]
    else:
        if ascending:
            sorted_index = df[by].argsort()
        else:
            sorted_index = df[by].argsort()[::-1]
    sorted_df = df.iloc[sorted_index].reset_index(drop=True)
    return sorted_df

def isa(*args,**kwargs):
    """
    fpath, contains='img'
    containss file paths based on the specified contains.
    Args:
        fpath (str): Path to the file.
        contains (str): contains of file to contains. Default is 'img' for images. Other options include 'doc' for documents,
                    'zip' for ZIP archives, and 'other' for other types of files.
    Returns:
        bool: True if the file matches the contains, False otherwise.
    """
    for arg in args:
        if isinstance(arg, str):
            if '/' in arg or '\\' in arg:
                fpath = arg
            else:
                contains=arg
    if 'img' in contains.lower() or 'image' in contains.lower():
        return is_image(fpath)
    elif 'doc' in contains.lower():
        return is_document(fpath)
    elif 'zip' in contains.lower():
        return is_zip(fpath)
    elif 'dir' in contains.lower() or ('f' in contains.lower() and 'd' in contains.lower()):
        return os.path.isdir(fpath)
    elif 'fi' in contains.lower():#file
        return os.path.isfile(fpath)
    elif 'num' in contains.lower():#file
        return os.path.isfile(fpath)
    elif 'text' in contains.lower() or 'txt' in contains.lower():#file
        return is_text(fpath)
    elif 'color' in contains.lower():#file
        return is_str_color(fpath)
    else:
        print(f"{contains} was not set up correctly")
        return False

def listdir(
    rootdir,
    kind="folder",
    sort_by="name",
    ascending=True,
    contains=None,
    orient="list",
    output="df"
):
    if not kind.startswith("."):
        kind = "." + kind

    if os.path.isdir(rootdir):
        ls = os.listdir(rootdir)
        fd = [".fd", ".fld", ".fol", ".fd", ".folder"]
        i = 0
        f = {
            "name": [],
            "length": [],
            "path": [],
            "created_time": [],
            "modified_time": [],
            "last_open_time": [],
            "size": [],
            "fname": [],
            "fpath": [],
        }
        for item in ls:
            item_path = os.path.join(rootdir, item)
            if item.startswith("."):
                continue
            filename, file_extension = os.path.splitext(item)
            is_folder = kind.lower() in fd and os.path.isdir(item_path)
            is_file = kind.lower() in file_extension.lower() and (
                os.path.isfile(item_path)
            )
            if kind in ['.doc','.img','.zip']: #选择大的类别
                if kind != ".folder" and not isa(item_path, kind):
                    continue
            elif kind in ['.all']:
                return flist(fpath, contains=contains)
            else: #精确到文件的后缀
                if not is_folder and not is_file:
                    continue
            f["name"].append(filename)
            f["length"].append(len(filename))
            f["path"].append(os.path.join(os.path.dirname(item_path), item))
            fpath = os.path.join(os.path.dirname(item_path), item)
            f["size"].append(round(os.path.getsize(fpath) / 1024 / 1024, 3))
            f["created_time"].append(pd.to_datetime(os.path.getctime(item_path), unit='s'))
            f["modified_time"].append(pd.to_datetime(os.path.getmtime(item_path), unit='s'))
            f['last_open_time'].append(pd.to_datetime(os.path.getatime(item_path), unit='s'))
            f["fname"].append(filename)  # will be removed
            f["fpath"].append(fpath)  # will be removed
            i += 1

        f["num"] = i
        f["rootdir"] = rootdir
        f["os"] = os.uname().machine
    else:
        raise FileNotFoundError(
            'The directory "{}" does NOT exist. Please check the directory "rootdir".'.format(
                rootdir
            )
        )

    f = pd.DataFrame(f)

    if contains is not None:
        f = f[f["name"].str.contains(contains, case=False)]

    if "nam" in sort_by.lower():
        f = sort_kind(f, by="name", ascending=ascending)
    elif "crea" in sort_by.lower():
        f = sort_kind(f, by="created_time", ascending=ascending)
    elif "modi" in sort_by.lower():
        f = sort_kind(f, by="modified_time", ascending=ascending)
    elif "s" in sort_by.lower() and "z" in sort_by.lower():
        f = sort_kind(f, by="size", ascending=ascending)

    if 'df' in output:
        return f
    else:
        if 'l' in orient.lower():  # list # default
            res_output = Box(f.to_dict(orient="list"))
            return res_output
        if 'd' in orient.lower():  # dict
            return Box(f.to_dict(orient="dict"))
        if 'r' in orient.lower():  # records
            return Box(f.to_dict(orient="records"))
        if 'in' in orient.lower():  # records
            return Box(f.to_dict(orient="index"))
        if 'se' in orient.lower():  # records
            return Box(f.to_dict(orient="series"))

# Example usage:
# result = listdir('your_root_directory')
# print(result)
# df=listdir("/", contains='sss',sort_by='name',ascending=False)
# print(df.fname.to_list(),"\n",df.fpath.to_list())  
def list_func(lib_name, opt="call"):
    if opt == "call":
        funcs = [func for func in dir(lib_name) if callable(getattr(lib_name, func))]
    else:
        funcs = dir(lib_name)
    return funcs
def func_list(lib_name, opt="call"):
    return list_func(lib_name, opt=opt)

def newfolder(*args, **kwargs):
    """
    newfolder(pardir, chdir)
    Args:
        pardir (dir): parent dir
        chdir (str): children dir
        overwrite (bool): overwrite?
    Returns:
        mkdir, giving a option if exists_ok or not
    """
    overwrite=kwargs.get("overwrite",False)
    for arg in args:
        if isinstance(arg, str):
            if "/" in arg or "\\" in arg:
                pardir=arg
                print(f'pardir{pardir}')
            else:
                chdir = arg
                print(f'chdir{chdir}')
        elif isinstance(arg,bool):
            overwrite=arg
            print(overwrite)
        else:
            print(f"{arg}Error: not support a {type(arg)} type")
    rootdir = []
    # Convert string to list
    if isinstance(chdir, str):
        chdir = [chdir]
    # Subfoldername should be unique
    chdir = list(set(chdir))
    if isinstance(pardir, str):  # Dir_parents should be 'str' type
        pardir = os.path.normpath(pardir)
    # Get the slash type: "/" or "\"
    stype = '/' if '/' in pardir else '\\'
    # Check if the parent directory exists and is a directory path
    if os.path.isdir(pardir):
        os.chdir(pardir)  # Set current path
        # Check if subdirectories are not empty
        if chdir:
            chdir.sort()
            # Create multiple subdirectories at once
            for folder in chdir:
                # Check if the subfolder already exists
                child_tmp = os.path.join(pardir, folder)
                if not os.path.isdir(child_tmp):
                    os.mkdir('./' + folder)
                    print(f'\n {folder} was created successfully!\n')
                else:
                    if overwrite:
                        shutil.rmtree(child_tmp)
                        os.mkdir('./' + folder)
                        print(f'\n {folder} overwrite! \n')
                    else:
                        print(f'\n {folder} already exists! \n')
                rootdir.append(child_tmp + stype)  # Note down
        else:
            print('\nWarning: Dir_child doesn\'t exist\n')
    else:
        print('\nWarning: Dir_parent is not a directory path\n')
    # Dir is the main output, if only one dir, then str type is inconvenient
    if len(rootdir) == 1:
        rootdir = rootdir[0]
    return rootdir

def figsave(*args,dpi=300):
    dir_save = None
    fname = None 
    for arg in args:
        if isinstance(arg, str):
            if '/' in arg or '\\' in arg:
                dir_save = arg
            elif '/' not in arg and '\\' not in arg:
                fname = arg
    # Backup original values
    if '/' in dir_save:
        if dir_save[-1] != '/':
            dir_save = dir_save + '/'
    elif '\\' in dir_save:
        if dir_save[-1] != '\\':
            dir_save = dir_save + '\\'
    else:
        raise ValueError('Check the Path of dir_save Directory')
    ftype = fname.split('.')[-1]
    if len(fname.split('.')) == 1:
        ftype = 'nofmt'
        fname = dir_save + fname + '.' + ftype
    else:
        fname = dir_save + fname
    # Save figure based on file type
    if ftype.lower() == 'eps':
        plt.savefig(fname, format='eps', bbox_inches='tight')
        plt.savefig(fname.replace('.eps', '.pdf'),
                    format='pdf', bbox_inches='tight',dpi=dpi)
    elif ftype.lower() == 'nofmt': # default: both "tif" and "pdf"
        fname_corr=fname.replace('nofmt','pdf')
        plt.savefig(fname_corr, format='pdf', bbox_inches='tight',dpi=dpi)
        fname=fname.replace('nofmt','tif')
        plt.savefig(fname, format='tiff', dpi=dpi, bbox_inches='tight')
        print(f"default saving filetype: both 'tif' and 'pdf")
    elif ftype.lower() == 'pdf':
        plt.savefig(fname, format='pdf', bbox_inches='tight',dpi=dpi)
    elif ftype.lower() in ['jpg', 'jpeg']:
        plt.savefig(fname, format='jpeg', dpi=dpi, bbox_inches='tight')
    elif ftype.lower() == 'png':
        plt.savefig(fname, format='png', dpi=dpi,
                    bbox_inches='tight', transparent=True)
    elif ftype.lower() in ['tiff', 'tif']:
        plt.savefig(fname, format='tiff', dpi=dpi, bbox_inches='tight')
    elif ftype.lower() == 'emf':
        plt.savefig(fname, format='emf', dpi=dpi, bbox_inches='tight')
    elif ftype.lower() == 'fig':
        plt.savefig(fname, format='pdf', bbox_inches='tight',dpi=dpi)
    print(f'\nSaved @: dpi={dpi}\n{fname}')


# ==============FuncStars(ax,x1=1,x2=2, yscale=0.9, pval=0.01)====================================================
# Usage:
# FuncStars(ax, x1=2, x2=3, yscale=0.99, pval=0.02)
# =============================================================================

# FuncStars --v 0.1.1
def FuncStars(ax,
              pval=None,
              Ylim=None,
              Xlim=None,
              symbol='*',
              yscale=0.95,
              x1=0,
              x2=1,
              alpha=0.05,
              fontsize=14,
              fontsize_note=6,
              rotation=0,
              fontname='Arial',
              values_below=None,
              linego=True,
              linestyle='-',
              linecolor='k',
              linewidth=.8,
              nsshow='off',
              symbolcolor='k',
              tailindicator=[0.06, 0.06],
              report=None,
              report_scale=-0.1,
              report_loc=None):
    if ax is None:
        ax = plt.gca()
    if Ylim is None:
        Ylim = plt.gca().get_ylim()
    if Xlim is None:
        Xlim = ax.get_xlim()
    if report_loc is None and report is not None:
        report_loc = np.min(Ylim) + report_scale*np.abs(np.diff(Ylim))
    if report_scale > 0:
        report_scale = -np.abs(report_scale)
    yscale = np.float64(yscale)
    y_loc = np.min(Ylim) + yscale*(np.max(Ylim)-np.min(Ylim))
    xcenter = np.mean([x1, x2])
    # ns / *
    if alpha < pval:
        if nsshow == 'on':
            ns_str = f'p={round(pval, 3)}' if pval < 0.9 else 'ns'
            color = 'm' if pval < 0.1 else 'k'
            plt.text(xcenter, y_loc, ns_str,
                     ha='center', va='bottom',  # 'center_baseline',
                     fontsize=fontsize-6 if fontsize > 6 else fontsize,
                     fontname=fontname, color=color, rotation=rotation
                     # bbox=dict(facecolor=None, edgecolor=None, color=None, linewidth=None)
                     )
    elif 0.01 < pval <= alpha:
        plt.text(xcenter, y_loc, symbol,
                 ha='center', va='center_baseline',
                 fontsize=fontsize, fontname=fontname, color=symbolcolor)
    elif 0.001 < pval <= 0.01:
        plt.text(xcenter, y_loc, symbol * 2,
                 ha='center', va='center_baseline',
                 fontsize=fontsize, fontname=fontname, color=symbolcolor)
    elif 0 < pval <= 0.001:
        plt.text(xcenter, y_loc, symbol * 3,
                 ha='center', va='center_baseline',
                 fontsize=fontsize, fontname=fontname, color=symbolcolor)
    # lines indicators
    if linego:  # and 0 < pval <= 0.05:
        print(pval)
        print(linego)
        # horizontal line
        if yscale < 0.99:
            plt.plot([x1 + np.abs(np.diff(Xlim)) * 0.01,
                      x2 - np.abs(np.diff(Xlim)) * 0.01],
                     [y_loc - np.abs(np.diff(Ylim)) * .03,
                      y_loc - np.abs(np.diff(Ylim)) * .03],
                     linestyle=linestyle, color=linecolor, linewidth=linewidth)
            # vertical line
            plt.plot([x1 + np.abs(np.diff(Xlim)) * 0.01,
                      x1 + np.abs(np.diff(Xlim)) * 0.01],
                     [y_loc - np.abs(np.diff(Ylim)) * tailindicator[0],
                      y_loc - np.abs(np.diff(Ylim)) * .03],
                     linestyle=linestyle, color=linecolor, linewidth=linewidth)
            plt.plot([x2 - np.abs(np.diff(Xlim)) * 0.01,
                      x2 - np.abs(np.diff(Xlim)) * 0.01],
                     [y_loc - np.abs(np.diff(Ylim)) * tailindicator[1],
                      y_loc - np.abs(np.diff(Ylim)) * .03],
                     linestyle=linestyle, color=linecolor, linewidth=linewidth)
        else:
            plt.plot([x1 + np.abs(np.diff(Xlim)) * 0.01,
                      x2 - np.abs(np.diff(Xlim)) * 0.01],
                     [np.min(Ylim) + 0.95*(np.max(Ylim)-np.min(Ylim)) - np.abs(np.diff(Ylim)) * 0.002,
                      np.min(Ylim) + 0.95*(np.max(Ylim)-np.min(Ylim)) - np.abs(np.diff(Ylim)) * 0.002],
                     linestyle=linestyle, color=linecolor, linewidth=linewidth)
            # vertical line
            plt.plot([x1 + np.abs(np.diff(Xlim)) * 0.01,
                      x1 + np.abs(np.diff(Xlim)) * 0.01],
                     [np.min(Ylim) + 0.95*(np.max(Ylim)-np.min(Ylim)) - np.abs(np.diff(Ylim)) * tailindicator[0],
                      np.min(Ylim) + 0.95*(np.max(Ylim)-np.min(Ylim)) - np.abs(np.diff(Ylim)) * 0.002],
                     linestyle=linestyle, color=linecolor, linewidth=linewidth)
            plt.plot([x2 - np.abs(np.diff(Xlim)) * 0.01,
                      x2 - np.abs(np.diff(Xlim)) * 0.01],
                     [np.min(Ylim) + 0.95*(np.max(Ylim)-np.min(Ylim)) - np.abs(np.diff(Ylim)) * tailindicator[1],
                      np.min(Ylim) + 0.95*(np.max(Ylim)-np.min(Ylim)) - np.abs(np.diff(Ylim)) * 0.002],
                     linestyle=linestyle, color=linecolor, linewidth=linewidth)
    if values_below is not None:
        plt.text(xcenter, y_loc * (-0.1), values_below,
                 ha='center', va='bottom',  # 'center_baseline', rotation=rotation,
                 fontsize=fontsize_note, fontname=fontname, color='k')
    # report / comments
    if report is not None:
        plt.text(xcenter, report_loc, report,
                 ha='left', va='bottom',  # 'center_baseline', rotation=rotation,
                 fontsize=fontsize_note, fontname=fontname, color='.7')
def is_str_color(s):
    # Regular expression pattern for hexadecimal color codes
    color_code_pattern = r"^#([A-Fa-f0-9]{6}|[A-Fa-f0-9]{8})$"
    return re.match(color_code_pattern, s) is not None
def is_num(s):
    """
    Check if a string can be converted to a number (int or float).
    Parameters:
    - s (str): The string to check.
    Returns:
    - bool: True if the string can be converted to a number, False otherwise.
    """
    try:
        float(s)  # Try converting the string to a float
        return True
    except ValueError:
        return False
def isnum(s):
    return is_num(s)
def is_image(fpath):
    mime_type, _ = mimetypes.guess_type(fpath)
    if mime_type and mime_type.startswith('image'):
        return True
    else:
        return False
def is_document(fpath):
    mime_type, _ = mimetypes.guess_type(fpath)
    if mime_type and (
        mime_type.startswith('text/') or
        mime_type == 'application/pdf' or
        mime_type == 'application/msword' or
        mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
        mime_type == 'application/vnd.ms-excel' or
        mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' or
        mime_type == 'application/vnd.ms-powerpoint' or
        mime_type == 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
    ):
        return True
    else:
        return False
def is_zip(fpath):
    mime_type, _ = mimetypes.guess_type(fpath)
    if mime_type == 'application/zip':
        return True
    else:
        return False
    
def stdshade(ax=None,*args, **kwargs):
    if (
        isinstance(ax, np.ndarray)
        and ax.ndim == 2
        and min(ax.shape) > 1
        and max(ax.shape) > 1
    ):
        y = ax
        ax = plt.gca()
    if ax is None:
        ax = plt.gca()
    alpha = 0.5
    acolor = "k"
    paraStdSem = "sem"
    plotStyle = "-"
    plotMarker = "none"
    smth = 1
    l_c_one = ["r", "g", "b", "m", "c", "y", "k", "w"]
    l_style2 = ["--", "-."]
    l_style1 = ["-", ":"]
    l_mark = ["o", "+", "*", ".", "x", "_", "|", "s", "d", "^", "v", ">", "<", "p", "h"]
    # Check each argument
    for iarg in range(len(args)):
        if (
            isinstance(args[iarg], np.ndarray)
            and args[iarg].ndim == 2
            and min(args[iarg].shape) > 1
            and max(args[iarg].shape) > 1
        ):
            y = args[iarg]
        # Except y, continuous data is 'F'
        if (isinstance(args[iarg], np.ndarray) and args[iarg].ndim == 1) or isinstance(
            args[iarg], range
        ):
            x = args[iarg]
            if isinstance(x, range):
                x = np.arange(start=x.start, stop=x.stop, step=x.step)
        # Only one number( 0~1), 'alpha' / color
        if isinstance(args[iarg], (int, float)):
            if np.size(args[iarg]) == 1 and 0 <= args[iarg] <= 1:
                alpha = args[iarg]
        if isinstance(args[iarg], (list, tuple)) and np.size(args[iarg]) == 3:
            acolor = args[iarg]
            acolor = tuple(acolor) if isinstance(acolor, list) else acolor
        # Color / plotStyle /
        if (
            isinstance(args[iarg], str)
            and len(args[iarg]) == 1
            and args[iarg] in l_c_one
        ):
            acolor = args[iarg]
        else:
            if isinstance(args[iarg], str):
                if args[iarg] in ["sem", "std"]:
                    paraStdSem = args[iarg]
                if args[iarg].startswith("#"):
                    acolor=hue2rgb(args[iarg])
                if str2list(args[iarg])[0] in l_c_one:
                    if len(args[iarg]) == 3:
                        k = [i for i in str2list(args[iarg]) if i in l_c_one]
                        if k != []:
                            acolor = k[0] 
                        st = [i for i in l_style2 if i in args[iarg]]
                        if st != []:
                            plotStyle = st[0] 
                    elif len(args[iarg]) == 2:
                        k = [i for i in str2list(args[iarg]) if i in l_c_one]
                        if k != []:
                            acolor = k[0] 
                        mk = [i for i in str2list(args[iarg]) if i in l_mark]
                        if mk != []:
                            plotMarker = mk[0] 
                        st = [i for i in l_style1 if i in args[iarg]]
                        if st != []:
                            plotStyle = st[0] 
                if len(args[iarg]) == 1:
                    k = [i for i in str2list(args[iarg]) if i in l_c_one]
                    if k != []:
                        acolor = k[0] 
                    mk = [i for i in str2list(args[iarg]) if i in l_mark]
                    if mk != []:
                        plotMarker = mk[0] 
                    st = [i for i in l_style1 if i in args[iarg]]
                    if st != []:
                        plotStyle = st[0] 
                if len(args[iarg]) == 2:
                    st = [i for i in l_style2 if i in args[iarg]]
                    if st != []:
                        plotStyle = st[0]
        # smth
        if (
            isinstance(args[iarg], (int, float))
            and np.size(args[iarg]) == 1
            and args[iarg] >= 1
        ):
            smth = args[iarg]

    if "x" not in locals() or x is None:
        x = np.arange(1, y.shape[1] + 1)
    elif len(x) < y.shape[1]:
        y = y[:, x]
        nRow = y.shape[0]
        nCol = y.shape[1]
        print(f"y was corrected, please confirm that {nRow} row, {nCol} col")
    else:
        x = np.arange(1, y.shape[1] + 1)

    if x.shape[0] != 1:
        x = x.T
    yMean = np.nanmean(y, axis=0)
    if smth > 1:
        yMean = savgol_filter(np.nanmean(y, axis=0), smth, 1)
    else:
        yMean = np.nanmean(y, axis=0)
    if paraStdSem == "sem":
        if smth > 1:
            wings = savgol_filter(np.nanstd(y, axis=0) / np.sqrt(y.shape[0]), smth, 1)
        else:
            wings = np.nanstd(y, axis=0) / np.sqrt(y.shape[0])
    elif paraStdSem == "std":
        if smth > 1:
            wings = savgol_filter(np.nanstd(y, axis=0), smth, 1)
        else:
            wings = np.nanstd(y, axis=0)

    fill_kws = kwargs.get('fill_kws', {})
    line_kws = kwargs.get('line_kws', {})
    fill = ax.fill_between(x, yMean + wings, yMean - wings, color=acolor, alpha=alpha, lw=0,**fill_kws)
    if line_kws != {} and not any(key.lower() in ['lw', 'linewidth'] for key in line_kws.keys()):
        line = ax.plot(x, yMean, color=acolor, lw=1.5, ls=plotStyle, marker=plotMarker, **line_kws)
    else:
        line = ax.plot(x, yMean, color=acolor, ls=plotStyle, marker=plotMarker, **line_kws)
    return line[0], fill
# =============================================================================
# # for plot figures {Qiu et al.2023}
# =============================================================================
# =============================================================================
# plt.rcParams.update({'figure.max_open_warning': 0})
# # Output matplotlib figure to SVG with text as text, not curves
# plt.rcParams['svg.fonttype'] = 'none'
# plt.rcParams['pdf.fonttype'] = 42
#
# plt.rc('text', usetex=False)
# # plt.style.use('ggplot')
# plt.style.use('science')
# plt.rc('font', family='serif')
# plt.rcParams.update({
#     "font.family": "serif",   # specify font family here
#     "font.serif": ["Arial"],  # specify font here
#     "font.size": 11})
# # plt.tight_layout()
# =============================================================================
# =============================================================================
# # axis spine
# # use it like: adjust_spines(ax, ['left', 'bottom'])
# =============================================================================


def adjust_spines(ax=None, spines=['left', 'bottom'],distance=2):
    if ax is None:
        ax = plt.gca()
    for loc, spine in ax.spines.items():
        if loc in spines:
            spine.set_position(('outward', distance))  # outward by 2 points
            # spine.set_smart_bounds(True)
        else:
            spine.set_color('none')  # don't draw spine
    # turn off ticks where there is no spine
    if 'left' in spines:
        ax.yaxis.set_ticks_position('left')
    else:
        ax.yaxis.set_ticks([])
    if 'bottom' in spines:
        ax.xaxis.set_ticks_position('bottom')
    else:
        # no xaxis ticks
        ax.xaxis.set_ticks([])
# And then plot the data:

def add_colorbar(im, width=None, pad=None, **kwargs):
    # usage: add_colorbar(im, width=0.01, pad=0.005, label="PSD (dB)", shrink=0.8)
    l, b, w, h = im.axes.get_position().bounds  # get boundaries
    width = width or 0.1 * w  # get width of the colorbar
    pad = pad or width  # get pad between im and cbar
    fig = im.axes.figure  # get figure of image
    cax = fig.add_axes([l + w + pad, b, width, h])  # define cbar Axes
    return fig.colorbar(im, cax=cax, **kwargs)  # draw cbar
# =============================================================================
# # for plot figures: setting rcParams
# usage: set_pub()
# or by setting sns.set_theme...see below:
# sns.set_theme(style="ticks", rc=params)      # 白色无刻度线，有坐标轴标度
# # sns.set_theme(style="whitegrid", rc=params)# 白色＋刻度线，无坐标轴标度
# # sns.set_theme(style="white", rc=params)    # 白色无刻度线，无坐标轴标度
# # sns.set_theme(style="dark", rc=params)     # 深色无刻度线，无坐标轴标度
# =============================================================================


def FuncCmpt(X1, X2, pmc='auto', pair='unpaired'):
    # output = {}

    # pmc correction: 'parametric'/'non-parametric'/'auto'
    # meawhile get the opposite setting (to compare the results)
    def corr_pmc(pmc):
        cfg_pmc = None
        if pmc.lower() in {'pmc', 'parametric'} and pmc.lower() not in {'npmc', 'nonparametric', 'non-parametric'}:
            cfg_pmc = 'parametric' 
        elif pmc.lower() in {'npmc', 'nonparametric', 'non-parametric'} and pmc.lower() not in {'pmc', 'parametric'}:
            cfg_pmc = 'non-parametric' 
        else:
            cfg_pmc = 'auto' 
        return cfg_pmc

    def corr_pair(pair):
        cfg_pair = None
        if 'pa' in pair.lower() and 'np' not in pair.lower():
            cfg_pair = 'paired'
        elif 'np' in pair.lower():
            cfg_pair = 'unpaired'
        return cfg_pair

    def check_normality(data):
        stat_shapiro, pval_shapiro = stats.shapiro(data)
        if pval_shapiro > 0.05:
            Normality = True
        else:
            Normality = False
        print(f'\n normally distributed\n') if Normality else print(
            f'\n NOT normally distributed\n')
        return Normality

    def sub_cmpt_2group(X1, X2, cfg_pmc='pmc', pair='unpaired'):
        output = {}
        nX1 = np.sum(~np.isnan(X1))
        nX2 = np.sum(~np.isnan(X2))
        if cfg_pmc == 'parametric' or cfg_pmc == 'auto':
            # VarType correction by checking variance Type via "levene"
            stat_lev, pval_lev = stats.levene(
                X1, X2, center='median', proportiontocut=0.05)
            VarType = True if pval_lev > 0.05 and nX1 == nX2 else False

            if 'np' in pair:  # 'unpaired'
                if VarType and Normality:
                    # The independent t-test requires that the dependent variable is approximately normally
                    # distributed within each group
                    # Note: Technically, it is the residuals that need to be normally distributed, but for
                    # an independent t-test, both will give you the same result.
                    stat_value, pval= stats.ttest_ind(
                        X1, X2, axis=0, equal_var=True, nan_policy='omit', alternative='two-sided')
                    notes_stat = 'unpaired t test'
                    notes_APA = f't({nX1+nX2-2})={round(stat_value, 5)},p={round(pval, 5)}'
                else:
                    # If the Levene's Test for Equality of Variances is statistically significant,
                    # which indicates that the group variances are unequal in the population, you
                    # can correct for this violation by not using the pooled estimate for the error
                    # term for the t-statistic, but instead using an adjustment to the degrees of
                    # freedom using the Welch-Satterthwaite method
                    stat_value, pval= stats.ttest_ind(
                        X1, X2, axis=0, equal_var=False, nan_policy='omit', alternative='two-sided')
                    notes_stat = 'Welchs t-test'
                    # note: APA FORMAT
                    notes_APA = f't({nX1+nX2-2})={round(stat_value, 5)},p={round(pval, 5)}'
            elif 'pa' in pair and 'np' not in pair:  # 'paired'
                # the paired-samples t-test is considered “robust” in handling violations of normality
                # to some extent. It can still yield valid results even if the data is not normally
                # distributed. Therefore, this test typically requires only approximately normal data
                stat_value, pval= stats.ttest_rel(
                    X1, X2, axis=0, nan_policy='omit', alternative='two-sided')
                notes_stat = 'paired t test'
                # note: APA FORMAT
                notes_APA = f't({sum([nX1-1])})={round(stat_value, 5)},p={round(pval, 5)}'
        elif cfg_pmc == 'non-parametric':
            if 'np' in pair:  # Perform Mann-Whitney
                stat_value, pval = stats.mannwhitneyu(
                    X1, X2, method='exact', nan_policy='omit')
                notes_stat = 'Mann-Whitney U'
                if nX1 == nX2:
                    notes_APA = f'U(n={nX1})={round(stat_value, 5)},p={round(pval, 5)}'
                else:
                    notes_APA = f'U(n1={nX1},n2={nX2})={round(stat_value, 5)},p={round(pval, 5)}'
            elif 'pa' in pair and 'np' not in pair:  # Wilcoxon signed-rank test
                stat_value, pval = stats.wilcoxon(
                    X1, X2, method='exact', nan_policy='omit')
                notes_stat = 'Wilcoxon signed-rank'
                if nX1 == nX2:
                    notes_APA = f'Z(n={nX1})={round(stat_value, 5)},p={round(pval, 5)}'
                else:
                    notes_APA = f'Z(n1={nX1},n2={nX2})={round(stat_value, 5)},p={round(pval, 5)}'

        # filling output
        output['stat'] = stat_value
        output['pval'] = pval
        output['method'] = notes_stat
        output['APA'] = notes_APA

        print(f"{output['method']}\n {notes_APA}\n\n")

        return output, pval

    Normality1 = check_normality(X1)
    Normality2 = check_normality(X2)
    Normality = True if all([Normality1, Normality2]) else False

    nX1 = np.sum(~np.isnan(X1))
    nX2 = np.sum(~np.isnan(X2))

    cfg_pmc = corr_pmc(pmc)
    cfg_pair = corr_pair(pair)

    output, p = sub_cmpt_2group(
        X1, X2, cfg_pmc=cfg_pmc, pair=cfg_pair)
    return p, output

# ======compare 2 group test===================================================
# # Example
# X1 = [19, 22, 16, 29, 24]
# X2 = [20, 11, 17, 12, 22]

# p, res= FuncCmpt(X1, X2, pmc='pmc', pair='unparrr')

# =============================================================================

# =============================================================================
# # method = ['anova',  # 'One-way and N-way ANOVA',
# #           'rm_anova',  # 'One-way and two-way repeated measures ANOVA',
# #           'mixed_anova',  # 'Two way mixed ANOVA',
# #           'welch_anova',  # 'One-way Welch ANOVA',
# #           'kruskal',  # 'Non-parametric one-way ANOVA'
# #           'friedman',  # Non-parametric one-way repeated measures ANOVA
# #           ]
# =============================================================================


# =============================================================================
# # method = ['anova',  # 'One-way and N-way ANOVA',
# #           'rm_anova',  # 'One-way and two-way repeated measures ANOVA',
# #           'mixed_anova',  # 'Two way mixed ANOVA',
# #           'welch_anova',  # 'One-way Welch ANOVA',
# #           'kruskal',  # 'Non-parametric one-way ANOVA'
# #           'friedman',  # Non-parametric one-way repeated measures ANOVA
# #           ]
# =============================================================================
def df_wide_long(df):
    rows, columns = df.shape 
    if columns > rows:
        return "Wide"
    elif rows > columns:
        return "Long"

def FuncMultiCmpt(pmc='pmc', pair='unpair', data=None, dv=None, factor=None,
                  ss_type=2, detailed=True, effsize='np2',
                  correction='auto', between=None, within=None,
                  subject=None, group=None
                  ):

    def corr_pair(pair):
        cfg_pair = None
        if 'pa' in pair.lower() and 'np' not in pair.lower():
            cfg_pair = 'paired'
        elif 'np' in pair.lower():
            cfg_pair = 'unpaired'
        elif 'mix' in pair.lower():
            cfg_pair = 'mix'
        return cfg_pair

    def check_normality(data):
        stat_shapiro, pval_shapiro = stats.shapiro(data)
        if pval_shapiro > 0.05:
            Normality = True
        else:
            Normality = False
        print(f'\n normally distributed\n') if Normality else print(
            f'\n NOT normally distributed\n')
        return Normality

    def corr_pmc(pmc):
        cfg_pmc = None
        if pmc.lower() in {'pmc', 'parametric'} and pmc.lower() not in {'upmc', 'npmc', 'nonparametric', 'non-parametric'}:
            cfg_pmc = 'parametric' 
        elif pmc.lower() in {'upmc', 'npmc', 'nonparametric', 'non-parametric'} and pmc.lower() not in {'pmc', 'parametric'}:
            cfg_pmc = 'non-parametric' 
        else:
            cfg_pmc = 'auto' 
        return cfg_pmc

    def extract_apa(res_tab):
        notes_APA = []
        if "ddof1" in res_tab:
            for irow in range(res_tab.shape[0]):
                note_tmp = f'{res_tab.Source[irow]}:F{round(res_tab.ddof1[irow]),round(res_tab.ddof2[irow])}={round(res_tab.F[irow], 5)},p={round(res_tab["p-unc"][irow], 5)}'
                notes_APA.append([note_tmp])
        elif "DF" in res_tab:
            print(res_tab.shape[0])
            for irow in range(res_tab.shape[0]-1):
                note_tmp = f'{res_tab.Source[irow]}:F{round(res_tab.DF[irow]),round(res_tab.DF[res_tab.shape[0]-1])}={round(res_tab.F[irow], 5)},p={round(res_tab["p-unc"][irow], 5)}'
                notes_APA.append([note_tmp])
            notes_APA.append(['NaN'])
        elif "DF1" in res_tab:  # in 'mix' case
            for irow in range(res_tab.shape[0]):
                note_tmp = f'{res_tab.Source[irow]}:F{round(res_tab.DF1[irow]),round(res_tab.DF2[irow])}={round(res_tab.F[irow], 5)},p={round(res_tab["p-unc"][irow], 5)}'
                notes_APA.append([note_tmp])
        return notes_APA

    def anovatable(res_tab):
        if 'df' in res_tab:  # statsmodels
            res_tab['mean_sq'] = res_tab[:]['sum_sq']/res_tab[:]['df']
            res_tab['est_sq'] = res_tab[:-1]['sum_sq'] / \
                sum(res_tab['sum_sq'])
            res_tab['omega_sq'] = (res_tab[:-1]['sum_sq']-(res_tab[:-1]['df'] *
                                                           res_tab['mean_sq'][-1]))/(sum(res_tab['sum_sq'])+res_tab['mean_sq'][-1])
        elif 'DF' in res_tab:
            res_tab['MS'] = res_tab[:]['SS']/res_tab[:]['DF']
            res_tab['est_sq'] = res_tab[:-1]['SS']/sum(res_tab['SS'])
            res_tab['omega_sq'] = (res_tab[:-1]['SS']-(res_tab[:-1]['DF'] *
                                                       res_tab['MS'][1]))/(sum(res_tab['SS'])+res_tab['MS'][1])
        if 'p-unc' in res_tab:
            if 'np2' in res_tab:
                res_tab['est_sq'] = res_tab['np2']
            if 'p-unc' in res_tab:
                res_tab['PR(>F)'] = res_tab['p-unc']
        return res_tab

    def run_anova(data, dv, factor, ss_type=2, detailed=True, effsize='np2'):
        # perform ANOVA
        # =============================================================================
        # #     # ANOVA (input: formula, dataset)
        # =============================================================================
        #     # note: if the data is balanced (equal sample size for each group), Type 1, 2, and 3 sums of squares
        #     # (typ parameter) will produce similar results.
        #     lm = ols("values ~ C(group)", data=df).fit()
        #     res_tab = anova_lm(lm, typ=ss_type)

        #     # however, it does not provide any effect size measures to tell if the
        #     # statistical significance is meaningful. The function below calculates
        #     # eta-squared () and omega-squared (). A quick note,  is the exact same
        #     # thing as  except when coming from the ANOVA framework people call it ;
        #     # is considered a better measure of effect size since it is unbiased in
        #     # it's calculation by accounting for the degrees of freedom in the model.
        #     # note: No effect sizes are calculated when using statsmodels.
        #     # to calculate eta squared, use the sum of squares from the table
        # res_tab = anovatable(res_tab)

        # =============================================================================
        #     # alternativ for ANOVA
        # =============================================================================
        res_tab = pg.anova(dv=dv, between=factor, data=data,
                           detailed=detailed, ss_type=ss_type, effsize=effsize)
        res_tab = anovatable(res_tab)
        return res_tab

    def run_rmanova(data, dv, factor, subject, correction='auto', detailed=True, effsize='ng2'):
        # One-way repeated-measures ANOVA using a long-format dataset.
        res_tab = pg.rm_anova(data=data, dv=dv, within=factor,
                              subject=subject, detailed=detailed, effsize=effsize)
        return res_tab

    def run_welchanova(data, dv, factor):
        # When the groups are balanced and have equal variances, the optimal
        # post-hoc test is the Tukey-HSD test (pingouin.pairwise_tukey()). If the
        # groups have unequal variances, the Games-Howell test is more adequate
        # (pingouin.pairwise_gameshowell()). Results have been tested against R.
        res_tab = pg.welch_anova(data=data, dv=dv, between=factor)
        res_tab = anovatable(res_tab)
        return res_tab

    def run_mixedanova(data, dv, between, within, subject, correction='auto', effsize='np2'):
        # Notes
        # Data are expected to be in long-format (even the repeated measures).
        # If your data is in wide-format, you can use the pandas.melt() function
        # to convert from wide to long format.

        # Warning
        # If the between-subject groups are unbalanced(=unequal sample sizes), a
        # type II ANOVA will be computed. Note however that SPSS, JAMOVI and JASP
        # by default return a type III ANOVA, which may lead to slightly different
        # results.
        res_tab = pg.mixed_anova(data=data, dv=dv, within=within, subject=subject,
                                 between=between, correction=correction, effsize=effsize)
        res_tab = anovatable(res_tab)
        return res_tab

    def run_friedman(data, dv, factor, subject, method='chisq'):
        # Friedman test for repeated measurements
        # The Friedman test is used for non-parametric (rank-based) one-way
        # repeated measures ANOVA

        # check df form ('long' or 'wide')
        # df_long = data.melt(ignore_index=False).reset_index()
        # if data.describe().shape[1] >= df_long.describe().shape[1]:
        #     res_tab = pg.friedman(data, method=method)
        # else:
        #     res_tab = pg.friedman(data=df_long, dv='value',
        #                           within="variable", subject="index", method=method)
        if "Wide" in df_wide_long(data):
            df_long = data.melt(ignore_index=False).reset_index()
            res_tab = pg.friedman(data=df_long, dv='value',
                                    within="variable", subject="index", method=method)
        else:
            res_tab = pg.friedman(data, dv=dv, within=factor, subject=subject,method=method)
        res_tab = anovatable(res_tab)
        return res_tab

    def run_kruskal(data, dv, factor):
        # Kruskal-Wallis H-test for independent samples
        res_tab = pg.kruskal(data=data, dv=dv, between=factor)
        res_tab = anovatable(res_tab)
        return res_tab

    # Normality Check:
    # Conduct normality tests (Shapiro-Wilk) for each group.
    # If the data is approximately normally distributed, ANOVA is robust to
    # moderate departures from normality, especially with larger sample sizes.

    # print(data[factor])
    # print(type(data[factor]))
    # print(len(data[factor].columns))
    # print(data[factor].nunique())
    # print(data[factor[0]])
    # print(data[factor[0]].unique())
    if group is None:
        group = factor

    # print(f'\ngroup is :\n{data[group]},\ndv is :\n{dv}\n')
    norm_array = []
    for sub_group in data[group].unique():
        norm_curr = check_normality(
            data.loc[data[group] == sub_group, dv])
        norm_array.append(norm_curr)
    norm_all = True if all(norm_array) else False

    # Homogeneity of Variances:
    # Check for homogeneity of variances (homoscedasticity) among groups.
    # Levene's test or Bartlett's test can be used for this purpose.
    # If variances are significantly different, consider transformations or use a
    # robust ANOVA method.

    # # =============================================================================
    # # # method1: stats.levene
    # # =============================================================================
    # # data_array = []
    # # for sub_group in df["group"].unique():
    # #     data_array.append(df.loc[df['group'] == sub_group, 'values'].values)
    # # print(data_array)
    # # variance_all = stats.levene(data_array[0],data_array[1],data_array[2])

    # =============================================================================
    # # method2: pingouin.homoscedasticity
    # =============================================================================
    res_levene = None
    variance_all = pg.homoscedasticity(
        data, dv=dv, group=group, method='levene', alpha=0.05)
    res_levene = True if variance_all.iloc[0,1] > 0.05 else False
    # =============================================================================
    # # ANOVA Assumptions:
    # # Ensure that the assumptions of independence, homogeneity of variances, and
    # # normality are reasonably met before proceeding.
    # =============================================================================
    notes_norm = 'normally' if norm_all else 'NOT-normally'
    notes_variance = 'equal' if res_levene else 'unequal'
    print(f'Data is {notes_norm} distributed, shows {notes_variance} variance')

    cfg_pmc = corr_pmc(pmc)
    cfg_pair = corr_pair(pair)
    output = {}
    if (cfg_pmc == 'parametric') or (cfg_pmc == 'auto'):
        if 'np' in cfg_pair:  # 'unpaired'
            if cfg_pmc == 'auto':
                if norm_all:
                    if res_levene:
                        res_tab = run_anova(data, dv, factor, ss_type=ss_type,
                                            detailed=True, effsize='np2')
                        notes_stat = f'{data[factor].nunique()} Way ANOVA'
                        notes_APA = extract_apa(res_tab)

                    else:
                        res_tab = run_welchanova(data, dv, factor)
                        notes_stat = f'{data[factor].nunique()} Way Welch ANOVA'
                        notes_APA = extract_apa(res_tab)

                else:

                    res_tab = run_kruskal(data, dv, factor)
                    notes_stat = f'Non-parametric Kruskal: {data[factor].nunique()} Way ANOVA'
                    notes_APA = extract_apa(res_tab)

            elif cfg_pmc == 'parametric':
                res_tab = run_anova(data, dv, factor, ss_type=ss_type,
                                    detailed=True, effsize='np2')
                notes_stat = f'{data[factor].nunique()} Way ANOVA'
                notes_APA = extract_apa(res_tab)

        elif 'pa' in cfg_pair and 'np' not in cfg_pair:  # 'paired'
            res_tab = run_rmanova(data, dv, factor, subject, correction='auto',
                                  detailed=True, effsize='ng2')
            notes_stat = f'{data[factor].nunique()} Way Repeated measures ANOVA'
            notes_APA = extract_apa(res_tab)

        elif 'mix' in cfg_pair or 'both' in cfg_pair:
            res_tab = run_mixedanova(data, dv, between, within, subject)
            # notes_stat = f'{len(sum(len(between)+sum(len(within))))} Way Mixed ANOVA'
            notes_stat = ""
            # n_inter = res_tab.loc(res_tab["Source"] == "Interaction")
            # print(n_inter)
            notes_APA = extract_apa(res_tab)

    elif cfg_pmc == 'non-parametric':
        if 'np' in cfg_pair:  # 'unpaired'
            res_tab = run_kruskal(data, dv, factor)
            notes_stat = f'Non-parametric Kruskal: {data[factor].nunique()} Way ANOVA'
            notes_APA = f'H({res_tab.ddof1[0]},n={data.shape[0]})={round(res_tab.H[0], 5)},p={round(res_tab["p-unc"][0], 5)}'

        elif 'pa' in cfg_pair and 'np' not in cfg_pair:  # 'paired'
            res_tab = run_friedman(data, dv, factor, subject, method='chisq')
            notes_stat = f'Non-parametric {data[factor].nunique()} Way Friedman repeated measures ANOVA'
            notes_APA = f'X^2({res_tab.ddof1[0]})={round(res_tab.Q[0], 5)},p={round(res_tab["p-unc"][0], 5)}'

    # =============================================================================
    # # Post-hoc
    # Post-Hoc Tests (if significant):
    # If ANOVA indicates significant differences, perform post-hoc tests (e.g.,
    # Tukey's HSD, Bonferroni, or Scheffé) to identify which groups differ from each other.
    # # https://pingouin-stats.org/build/html/generated/pingouin.pairwise_tests.html
    # =============================================================================
    go_pmc = True if cfg_pmc == 'parametric' else False
    go_subject = subject if ('pa' in cfg_pair) and (
        'np' not in cfg_pair) else None
    go_mix_between = between if ('mix' in cfg_pair) or (
        'both' in cfg_pair) else None
    go_mix_between = None if ('pa' in cfg_pair) or (
        'np' not in cfg_pair) else factor
    go_mix_within = within if ('mix' in cfg_pair) or (
        'both' in cfg_pair) else None
    go_mix_within = factor if ('pa' in cfg_pair) or (
        'np' not in cfg_pair) else None
    if res_tab['p-unc'][0] <= .05:
        # Pairwise Comparisons
        method_post_hoc = [
            "bonf",  # 'bonferroni',  # : one-step correction
            "sidak",  # one-step correction
            "holm",  # step-down method using Bonferroni adjustments
            "fdr_bh",  # Benjamini/Hochberg (non-negative)
            "fdr_by",  # Benjamini/Yekutieli (negative)
        ]
        res_posthoc = pd.DataFrame()
        for met in method_post_hoc:
            post_curr = pg.pairwise_tests(data=data, dv=dv, between=go_mix_between, within=go_mix_within, subject=go_subject, parametric=go_pmc, marginal=True, alpha=0.05, alternative='two-sided',
                                          padjust=met)
            res_posthoc = pd.concat([res_posthoc, post_curr],
                                    ignore_index=True)
    else:
        res_posthoc = None
    output['res_posthoc'] = res_posthoc
    # =============================================================================
    #     # filling output
    # =============================================================================

    pd.set_option('display.max_columns', None)
    output['stat'] = notes_stat
    # print(output['APA'])
    output['APA'] = notes_APA
    output['pval'] = res_tab['p-unc']
    output['res_tab'] = res_tab
    if res_tab.shape[0] == len(notes_APA):
        output['res_tab']['APA'] = output['APA']  # note APA in the table
    # print(output['stat'])
    # print(output['res_tab'])
    return output


# =============================================================================
# # One-way ANOVA
# =============================================================================
# url = "http://stats191.stanford.edu/data/rehab.csv"
# rehab_table = pd.read_table(url, delimiter=",")
# rehab_table.to_csv("rehab.table")
# fig, ax = plt.subplots(figsize=(8, 6))
# fig = rehab_table.boxplot("Time", "Fitness", ax=ax, grid=False)
# # fig, ax = plt.subplots(figsize=(8, 6))
# # set_pub()
# # sns.boxenplot(x="Time",y="Fitness",data = rehab_table)

# out2 = FuncMultiCmpt(pmc='pmc', pair='unpair',
#                       data=rehab_table, dv='Time', factor='Fitness')
# # print(out2['res_tab'])
# # print(out2['APA'])
# out2['res_posthoc']
# out2['res_posthoc']['p-unc'][0]
# out2['res_posthoc']['p-adjust'][0]
# out2['res_posthoc']['p-corr'][0]


# =============================================================================
# # Interactions and ANOVA
# https://www.statsmodels.org/dev/examples/notebooks/generated/interactions_anova.html
# url = "http://stats191.stanford.edu/data/salary.table"
# fh = urlopen(url)
# df = pd.read_table(fh)
# out1 = FuncMultiCmpt(pmc='pmc', pair='unpaired', data=df,
#                      dv='S', factor=['X', 'E', 'M'], group='M')
# # # two-way anova
# # https://www.statology.org/two-way-anova-python/
# # =============================================================================
# # df = pd.DataFrame({'water': np.repeat(['daily', 'weekly'], 15),
# #                    'sun': np.tile(np.repeat(['low', 'med', 'high'], 5), 2),
# #                    'height': [6, 6, 6, 5, 6, 5, 5, 6, 4, 5,
# #                               6, 6, 7, 8, 7, 3, 4, 4, 4, 5,
# #                               4, 4, 4, 4, 4, 5, 6, 6, 7, 8]})
# # out1 = FuncMultiCmpt(pmc='pmc', pair='unpaired', data=df,
# #                       dv='height', factor=['water','sun'],group='water')


# =============================================================================
# # two way anova
# https://www.geeksforgeeks.org/how-to-perform-a-two-way-anova-in-python/
# =============================================================================
# df1=pd.DataFrame({'Fertilizer': np.repeat(['daily', 'weekly'], 15),
#                           'Watering': np.repeat(['daily', 'weekly'], 15),
#                           'height': [14, 16, 15, 15, 16, 13, 12, 11,
#                                       14, 15, 16, 16, 17, 18, 14, 13,
#                                       14, 14, 14, 15, 16, 16, 17, 18,
#                                       14, 13, 14, 14, 14, 15]})

# df1['subject'] = np.tile(range(0, 15), (1, 2)).T
# out1 = FuncMultiCmpt(pmc='pmc', pair='unpaired', data=df1,
#                       dv='height', factor=['Fertilizer','Watering'],group='Watering')
# # print(out1['stat'])
# # print(out1['res_tab'])

# =============================================================================
# # welch anova
# https://www.geeksforgeeks.org/how-to-perform-welchs-anova-in-python/
# =============================================================================
# df = pd.DataFrame({'score': [64, 66, 68, 75, 78, 94, 98, 79, 71, 80,
#                              91, 92, 93, 90, 97, 94, 82, 88, 95, 96,
#                              79, 78, 88, 94, 92, 85, 83, 85, 82, 81],
#                    'group': np.repeat(['strat1', 'strat2', 'strat3'],repeats=10)})
# out1 = FuncMultiCmpt(pmc='auto',pair='unpaired',data=df, dv='score', factor='group', group='group')
# =============================================================================
# # two way anova
# https://www.statology.org/two-way-anova-python/
# =============================================================================
# df = pd.DataFrame({'water': np.repeat(['daily', 'weekly'], 15),
#                    'sun': np.tile(np.repeat(['low', 'med', 'high'], 5), 2),
#                    'height': [6, 6, 6, 5, 6, 5, 5, 6, 4, 5,
#                               6, 6, 7, 8, 7, 3, 4, 4, 4, 5,
#                               4, 4, 4, 4, 4, 5, 6, 6, 7, 8]})
# df['subject'] = np.tile(range(0, 15), (1, 2)).T
# out1 = FuncMultiCmpt(pmc='pmc', pair='unpaired', data=df,
#                      dv='height', factor=['water', 'sun'], subject='subject', group='water')
# # print(out1['stat'])
# # print(out1['res_tab'])

# =============================================================================
# # 3-way ANOVA
# =============================================================================
# df = pd.DataFrame({'program': np.repeat([1, 2], 20),
#                    'gender': np.tile(np.repeat(['M', 'F'], 10), 2),
#                    'division': np.tile(np.repeat([1, 2], 5), 4),
#                    'height': [7, 7, 8, 8, 7, 6, 6, 5, 6, 5,
#                               5, 5, 4, 5, 4, 3, 3, 4, 3, 3,
#                               6, 6, 5, 4, 5, 4, 5, 4, 4, 3,
#                               2, 2, 1, 4, 4, 2, 1, 1, 2, 1]})
# df['subject'] = np.tile(range(0, 20), (1, 2)).T
# out1 = FuncMultiCmpt(pmc='pmc', pair='unpaired', data=df,
#                      dv='height', factor=['gender', 'program', 'division'], subject='subject', group='program')
# # print(out1['stat'])
# # print(out1['res_tab'])

# =============================================================================
# # Repeated Measures ANOVA in Python
# =============================================================================
# df = pd.DataFrame({'patient': np.repeat([1, 2, 3, 4, 5], 4),
#                     'drug': np.tile([1, 2, 3, 4], 5),
#                     'response': [30, 28, 16, 34,
#                                 14, 18, 10, 22,
#                                 24, 20, 18, 30,
#                                 38, 34, 20, 44,
#                                 26, 28, 14, 30]})
# # df['subject'] = np.tile(range(0, 20), (1, 2)).T
# out1 = FuncMultiCmpt(pmc='pmc', pair='paired', data=df,
#                       dv='response', factor=['drug'], subject='patient', group='drug')
# print(out1['stat'])
# print(out1['res_tab'])
# print(out1['APA'])

# =============================================================================
# # repeated anova
# https://www.geeksforgeeks.org/how-to-perform-a-repeated-measures-anova-in-python/
# =============================================================================
# df = pd.DataFrame({'Cars': np.repeat([1, 2, 3, 4, 5], 4),
#                'Engine Oil': np.tile([1, 2, 3, 4], 5),
#                'Mileage': [36, 38, 30, 29,
#                            34, 38, 30, 29,
#                            34, 28, 38, 32,
#                            38, 34, 20, 44,
#                            26, 28, 34, 50]})
# out1 = FuncMultiCmpt(pmc='pmc', pair='paired', data=df,
#                  dv='Mileage', factor=['Engine Oil'], subject='Cars', group='Cars')
# =============================================================================
# #two-way repeated anova
# =============================================================================
# df = pd.read_csv(
#     "https://reneshbedre.github.io/assets/posts/anova/plants_leaves_two_within.csv")
# df
# # df['subject'] = np.tile(range(0, 20), (1, 2)).T
# out1 = FuncMultiCmpt(pmc='pmc', pair='paired', data=df,
#                       dv='num_leaves', factor=['year', 'time'], subject='plants', group='year')
# print(out1['stat'])
# print(out1['res_tab'])
# print(out1['APA'])

# =============================================================================
# # repeated anova
# =============================================================================
# df = pd.read_csv('/Users/macjianfeng/Desktop/test.csv')
# df.head()
# df.loc[df['animal'].str.contains('Sleep'), 'experiment'] = 'sleep'
# df.loc[df['animal'].str.contains('Wake'), 'experiment'] = 'wake'
# df.loc[df['variable'].str.contains('hypo'), 'region'] = 'hypo'
# df.loc[df['variable'].str.contains('cort'), 'region'] = 'cort'
# df
# for i in range(4):
#     match i:
#         case 0:
#             prot_name = 'A1'
#         case 1:
#             prot_name = 'A2'
#         case 2:
#             prot_name = '845'
#         case 3:
#             prot_name = '831'
#     df_tmp = df[df["variable"].str.contains(prot_name)]
#     df_tmp['protein'] = prot_name
#     df_tmp = df_tmp.reset_index()
#     print(df_tmp)

# out1 = FuncMultiCmpt(pmc='pmc', pair='mix', data=df_tmp,
#                      dv='value', between='experiment', within='region', subject='animal', group='experiment')
# print(out1['stat'])
# print(out1['res_tab'])
# # =============================================================================
# One-way ANOVA
# df1 = pd.read_csv('/Users/macjianfeng/Desktop/Book2.csv')
# df2 = df1.melt()
# out1 = FuncMultiCmpt(pmc='npmc', pair='unpaired', data=df2,
#                      dv='libido', factor=['brand x', 'brand y', 'brand z'], subject='participant')
# print(out1['stat'])
# print(out1['res_tab'])
# =============================================================================


# =============================================================================
# # #One-way ANOVA new example: https://www.pythonfordatascience.org/anova-python/
# =============================================================================
# df1 = pd.read_csv(
#     "https://raw.githubusercontent.com/researchpy/Data-sets/master/difficile.csv")
# df1.drop('person', axis=1, inplace=True)
# # Recoding value from numeric to string
# df1['dose'].replace({1: 'placebo', 2: 'low', 3: 'high'}, inplace=True)
# df1.head(10)

# out3= FuncMultiCmpt(pmc='pmc', data=df1, dv='libido', factor='dose')
# # print(out3['res_tab'])
# # # print(out3['res_posthoc'])
# # print(out3['APA'])

# =============================================================================
# https://lifewithdata.com/2023/06/08/how-to-perform-a-two-way-anova-in-python/
# =============================================================================
# data = {
#     'Diet': ['A', 'A', 'A', 'A', 'B', 'B', 'B', 'B', 'C', 'C', 'C', 'C'],
#     'Workout': ['Low', 'Medium', 'High', 'Low', 'Medium', 'High', 'Low', 'Medium', 'High', 'Low', 'Medium', 'High'],
#     'WeightLoss': [3, 4, 5, 3.2, 5, 6, 5.2, 6, 5.5, 4, 5.5, 6.2]
# }
# df = pd.DataFrame(data)
# out4= FuncMultiCmpt(pmc='pmc', pair='unpaired',data=df, dv='WeightLoss', factor=['Diet','Workout'],group='Diet')

# =============================================================================
# # convert to list to string
# =============================================================================
def list2str(x_str):
    s = ''.join(str(x) for x in x_str)
    return s
def str2list(str_):
    l = []
    [l.append(x) for x in str_]
    return l

def load_img(fpath):
    """
    Load an image from the specified file path.
    Args:
        fpath (str): The file path to the image.
    Returns:
        PIL.Image: The loaded image.
    Raises:
        FileNotFoundError: If the specified file is not found.
        OSError: If the specified file cannot be opened or is not a valid image file.
    """
    try:
        img = Image.open(fpath)
        return img
    except FileNotFoundError:
        raise FileNotFoundError(f"The file '{fpath}' was not found.")
    except OSError:
        raise OSError(f"Unable to open file '{fpath}' or it is not a valid image file.")

def apply_filter(img, *args):
    # def apply_filter(img, filter_name, filter_value=None):
    """
    Apply the specified filter to the image.
    Args:
        img (PIL.Image): The input image.
        filter_name (str): The name of the filter to apply.
        **kwargs: Additional parameters specific to the filter.
    Returns:
        PIL.Image: The filtered image.
    """
    def correct_filter_name(filter_name):
        if 'bl' in filter_name.lower() and 'box' not in filter_name.lower():
            return 'BLUR'
        elif 'cont' in  filter_name.lower():
            return 'Contour'
        elif 'det' in  filter_name.lower():
            return 'Detail'
        elif 'edg' in  filter_name.lower() and 'mo' not in filter_name.lower() and 'f' not in filter_name.lower():
            return 'EDGE_ENHANCE'
        elif 'edg' in  filter_name.lower() and 'mo' in filter_name.lower():
            return 'EDGE_ENHANCE_MORE'
        elif 'emb' in  filter_name.lower():
            return 'EMBOSS'
        elif 'edg' in  filter_name.lower() and 'f' in filter_name.lower():
            return 'FIND_EDGES'
        elif 'sh' in  filter_name.lower() and 'mo' not in filter_name.lower():
            return 'SHARPEN'
        elif 'sm' in  filter_name.lower() and 'mo' not in filter_name.lower():
            return 'SMOOTH'
        elif 'sm' in  filter_name.lower() and 'mo' in filter_name.lower():
            return 'SMOOTH_MORE'
        elif 'min' in  filter_name.lower():
            return 'MIN_FILTER'
        elif 'max' in  filter_name.lower():
            return 'MAX_FILTER'
        elif 'mod' in  filter_name.lower():
            return 'MODE_FILTER'
        elif 'mul' in  filter_name.lower():
            return 'MULTIBAND_FILTER'
        elif 'gau' in  filter_name.lower():
            return 'GAUSSIAN_BLUR'
        elif 'box' in  filter_name.lower():
            return 'BOX_BLUR'
        elif 'med' in  filter_name.lower():
            return 'MEDIAN_FILTER'
        else: 
            supported_filters = [
                "BLUR",
                "CONTOUR",
                "DETAIL",
                "EDGE_ENHANCE",
                "EDGE_ENHANCE_MORE",
                "EMBOSS",
                "FIND_EDGES",
                "SHARPEN",
                "SMOOTH",
                "SMOOTH_MORE",
                "MIN_FILTER",
                "MAX_FILTER",
                "MODE_FILTER",
                "MULTIBAND_FILTER",
                "GAUSSIAN_BLUR",
                "BOX_BLUR",
                "MEDIAN_FILTER",
            ]
            raise ValueError(
                f"Unsupported filter: {filter_name}, should be one of: {supported_filters}"
            )

    for arg in args:
        if isinstance(arg, str):
            filter_name = arg
            filter_name = correct_filter_name(filter_name)
        else:
            filter_value = arg
    filter_name = filter_name.upper()  # Ensure filter name is uppercase

    # Supported filters
    supported_filters = {
        "BLUR": ImageFilter.BLUR,
        "CONTOUR": ImageFilter.CONTOUR,
        "DETAIL": ImageFilter.DETAIL,
        "EDGE_ENHANCE": ImageFilter.EDGE_ENHANCE,
        "EDGE_ENHANCE_MORE": ImageFilter.EDGE_ENHANCE_MORE,
        "EMBOSS": ImageFilter.EMBOSS,
        "FIND_EDGES": ImageFilter.FIND_EDGES,
        "SHARPEN": ImageFilter.SHARPEN,
        "SMOOTH": ImageFilter.SMOOTH,
        "SMOOTH_MORE": ImageFilter.SMOOTH_MORE,
        "MIN_FILTER": ImageFilter.MinFilter,
        "MAX_FILTER": ImageFilter.MaxFilter,
        "MODE_FILTER": ImageFilter.ModeFilter,
        "MULTIBAND_FILTER": ImageFilter.MultibandFilter,
        "GAUSSIAN_BLUR": ImageFilter.GaussianBlur,
        "BOX_BLUR": ImageFilter.BoxBlur,
        "MEDIAN_FILTER": ImageFilter.MedianFilter,
    }
    # Check if the filter name is supported
    if filter_name not in supported_filters:
        raise ValueError(
            f"Unsupported filter: {filter_name}, should be one of: {[i.lower() for i in supported_filters.keys()]}"
        )

    # Apply the filter
    if filter_name.upper() in [
        "BOX_BLUR",
        "GAUSSIAN_BLUR",
        "MEDIAN_FILTER",
        "MIN_FILTER",
        "MAX_FILTER",
        "MODE_FILTER",
    ]:
        radius = filter_value if filter_value is not None else 2
        return img.filter(supported_filters[filter_name](radius))
    elif filter_name in ["MULTIBAND_FILTER"]:
        bands = filter_value if filter_value is not None else None
        return img.filter(supported_filters[filter_name](bands))
    else:
        if filter_value is not None:
            print(f"{filter_name} doesn't require a value for {filter_value}, but it remains unaffected")
        return img.filter(supported_filters[filter_name])


def imgsets(
    img,
    sets=None,
    show=True,
    show_axis=False,
    size=None,
    dpi=100,
    figsize=None,
    auto=False,
    filter_kws=None,
):
    """
    Apply various enhancements and filters to an image using PIL's ImageEnhance and ImageFilter modules.

    Args:
        img (PIL.Image): The input image.
        sets (dict): A dictionary specifying the enhancements, filters, and their parameters.
        show (bool): Whether to display the enhanced image.
        show_axis (bool): Whether to display axes on the image plot.
        size (tuple): The size of the thumbnail, cover, contain, or fit operation.
        dpi (int): Dots per inch for the displayed image.
        figsize (tuple): The size of the figure for displaying the image.
        auto (bool): Whether to automatically enhance the image based on its characteristics.

    Returns:
        PIL.Image: The enhanced image.

    Supported enhancements and filters:
        - "sharpness": Adjusts the sharpness of the image. Values > 1 increase sharpness, while values < 1 decrease sharpness.
        - "contrast": Adjusts the contrast of the image. Values > 1 increase contrast, while values < 1 decrease contrast.
        - "brightness": Adjusts the brightness of the image. Values > 1 increase brightness, while values < 1 decrease brightness.
        - "color": Adjusts the color saturation of the image. Values > 1 increase saturation, while values < 1 decrease saturation.
        - "rotate": Rotates the image by the specified angle.
        - "crop" or "cut": Crops the image. The value should be a tuple specifying the crop box as (left, upper, right, lower).
        - "size": Resizes the image to the specified dimensions.
        - "thumbnail": Resizes the image to fit within the given size while preserving aspect ratio.
        - "cover": Resizes and crops the image to fill the specified size.
        - "contain": Resizes the image to fit within the specified size, adding borders if necessary.
        - "fit": Resizes and pads the image to fit within the specified size.
        - "filter": Applies various filters to the image (e.g., BLUR, CONTOUR, EDGE_ENHANCE).

    Note:
        The "color" and "enhance" enhancements are not implemented in this function.
    """
    supported_filters = [
                "BLUR",
                "CONTOUR",
                "DETAIL",
                "EDGE_ENHANCE",
                "EDGE_ENHANCE_MORE",
                "EMBOSS",
                "FIND_EDGES",
                "SHARPEN",
                "SMOOTH",
                "SMOOTH_MORE",
                "MIN_FILTER",
                "MAX_FILTER",
                "MODE_FILTER",
                "MULTIBAND_FILTER",
                "GAUSSIAN_BLUR",
                "BOX_BLUR",
                "MEDIAN_FILTER",
            ]
    print("sets: a dict,'sharp:1.2','color','contrast:'auto' or 1.2','bright', 'crop: x_upperleft,y_upperleft, x_lowerright, y_lowerright','rotation','resize','rem or background'")
    print(f"usage: filter_kws 'dict' below:")
    pp([str(i).lower() for i in supported_filters])
    print("\nlog:\n")
    def confirm_rembg_models(model_name):
        models_support = [
            "u2net",
            "u2netp",
            "u2net_human_seg",
            "u2net_cloth_seg",
            "silueta",
            "isnet-general-use",
            "isnet-anime",
            "sam",
        ]
        if model_name in models_support:
            print(f"model_name: {model_name}")
            return model_name
        else:
            print(f"{model_name} cannot be found, check the name:{models_support}, default('isnet-general-use') has been used")
            return "isnet-general-use"
    def auto_enhance(img):
        """
        Automatically enhances the image based on its characteristics.
        Args:
            img (PIL.Image): The input image.
        Returns:
            dict: A dictionary containing the optimal enhancement values.
        """
        # Determine the bit depth based on the image mode
        if img.mode in ["1", "L", "P", "RGB", "YCbCr", "LAB", "HSV"]:
            # 8-bit depth per channel
            bit_depth = 8
        elif img.mode in ["RGBA", "CMYK"]:
            # 8-bit depth per channel + alpha (RGBA) or additional channels (CMYK)
            bit_depth = 8
        elif img.mode in ["I", "F"]:
            # 16-bit depth per channel (integer or floating-point)
            bit_depth = 16
        else:
            raise ValueError("Unsupported image mode")
        # Calculate the brightness and contrast for each channel
        num_channels = len(img.getbands())
        brightness_factors = []
        contrast_factors = []
        for channel in range(num_channels):
            channel_histogram = img.split()[channel].histogram()
            brightness = sum(i * w for i, w in enumerate(channel_histogram))/sum(channel_histogram)
            channel_min, channel_max = img.split()[channel].getextrema()
            contrast = channel_max - channel_min
            # Adjust calculations based on bit depth
            normalization_factor = 2**bit_depth - 1  # Max value for the given bit depth
            brightness_factor = (1.0 + (brightness - normalization_factor / 2) / normalization_factor)
            contrast_factor = (1.0 + (contrast - normalization_factor / 2) / normalization_factor)
            brightness_factors.append(brightness_factor)
            contrast_factors.append(contrast_factor)
        # Calculate the average brightness and contrast factors across channels
        avg_brightness_factor = sum(brightness_factors) / num_channels
        avg_contrast_factor = sum(contrast_factors) / num_channels
        return {"brightness": avg_brightness_factor, "contrast": avg_contrast_factor}
    # Load image if input is a file path
    if isinstance(img, str):
        img = load_img(img)
    img_update = img.copy()
    # Auto-enhance image if requested
    if auto:
        auto_params = auto_enhance(img_update)
        sets.update(auto_params)
    if sets is None:
        sets = {}
    for k, value in sets.items():
        if "shar" in k.lower():
            enhancer = ImageEnhance.Sharpness(img_update)
            img_update = enhancer.enhance(value)
        elif "col" in k.lower() and 'bg' not in k.lower():
            enhancer = ImageEnhance.Color(img_update)
            img_update = enhancer.enhance(value)
        elif "contr" in k.lower():
            if value and isinstance(value,(float,int)):
                enhancer = ImageEnhance.Contrast(img_update)
                img_update = enhancer.enhance(value)
            else:
                print('autocontrasted')
                img_update = ImageOps.autocontrast(img_update)
        elif "bri" in k.lower():
            enhancer = ImageEnhance.Brightness(img_update)
            img_update = enhancer.enhance(value)
        elif "cro" in k.lower() or "cut" in k.lower():
            img_update=img_update.crop(value)
        elif "rota" in k.lower():
            img_update = img_update.rotate(value)
        elif "si" in k.lower():
            img_update = img_update.resize(value)
        elif "thum" in k.lower():
            img_update.thumbnail(value)
        elif "cover" in k.lower():
            img_update = ImageOps.cover(img_update, size=value)
        elif "contain" in k.lower():
            img_update = ImageOps.contain(img_update, size=value)
        elif "fit" in k.lower():
            img_update = ImageOps.fit(img_update, size=value)
        elif "pad" in k.lower():
            img_update = ImageOps.pad(img_update, size=value)
        elif 'rem' in k.lower() or 'rm' in k.lower() or 'back' in k.lower():
            if value and isinstance(value,(int,float,list)):
                print('example usage: {"rm":[alpha_matting_background_threshold(20),alpha_matting_foreground_threshold(270),alpha_matting_erode_sive(11)]}')
                print("https://github.com/danielgatis/rembg/blob/main/USAGE.md")
                #     ###            Parameters:
                #         data (Union[bytes, PILImage, np.ndarray]): The input image data.
                #         alpha_matting (bool, optional): Flag indicating whether to use alpha matting. Defaults to False.
                #         alpha_matting_foreground_threshold (int, optional): Foreground threshold for alpha matting. Defaults to 240.
                #         alpha_matting_background_threshold (int, optional): Background threshold for alpha matting. Defaults to 10.
                #         alpha_matting_erode_size (int, optional): Erosion size for alpha matting. Defaults to 10.
                #         session (Optional[BaseSession], optional): A session object for the 'u2net' model. Defaults to None.
                #         only_mask (bool, optional): Flag indicating whether to return only the binary masks. Defaults to False.
                #         post_process_mask (bool, optional): Flag indicating whether to post-process the masks. Defaults to False.
                #         bgcolor (Optional[Tuple[int, int, int, int]], optional): Background color for the cutout image. Defaults to None.
                #  ###
                if isinstance(value,int):
                    value=[value]
                if len(value) <2:
                    img_update = remove(img_update,alpha_matting=True,alpha_matting_background_threshold=value)
                elif 2<=len(value)<3:
                    img_update = remove(img_update,alpha_matting=True,alpha_matting_background_threshold=value[0],alpha_matting_foreground_threshold=value[1])
                elif 3<=len(value)<4:
                    img_update = remove(img_update,alpha_matting=True,alpha_matting_background_threshold=value[0],alpha_matting_foreground_threshold=value[1],alpha_matting_erode_size=value[2])
            if isinstance(value,tuple): # replace the background color
                if len(value)==3:
                    value+=(255,)
                img_update = remove(img_update, bgcolor=value)
            if isinstance(value,str):
                if confirm_rembg_models(value):
                    img_update=remove(img_update,session=new_session(value))
                else:
                    img_update=remove(img_update)
        elif 'bgcolor' in k.lower():
            if isinstance(value,list):
                value=tuple(value)
            if isinstance(value,tuple): # replace the background color
                if len(value)==3:
                    value+=(255,)
                img_update = remove(img_update, bgcolor=value)
    if filter_kws:
        for filter_name, filter_value in filter_kws.items():
            img_update = apply_filter(img_update, filter_name, filter_value)
    # Display the image if requested
    if show:
        if figsize is None:
            plt.figure(dpi=dpi)
        else:
            plt.figure(figsize=figsize, dpi=dpi)
        plt.imshow(img_update)
        plt.axis("on") if show_axis else plt.axis("off")
    return img_update
# # usage:
# img = imgsets(
#     fpath,
#     sets={"rota": -5},
#     dpi=200,
#     filter_kws={"EMBOSS": 5, "sharpen": 5, "EDGE_ENHANCE_MORE": 10},
#     show_axis=True,
# )

def figsets(*args):
    fig = plt.gcf()
    fontsize = 11
    fontname = "Arial"
    sns_themes = ["white", "whitegrid", "dark", "darkgrid", "ticks"]
    sns_contexts = ["notebook", "talk", "poster"]  # now available "paper"
    scienceplots_styles = ["science","nature",
        "scatter","ieee","no-latex","std-colors","high-vis","bright","dark_background","science",
        "high-vis","vibrant","muted","retro","grid","high-contrast","light","cjk-tc-font","cjk-kr-font",
    ]
    def sets_priority(ax,key, value):
        if ("fo" in key) and (("size" in key) or ("sz" in key)):
            fontsize=value
            plt.rcParams.update({"font.size": value})
        # style
        if "st" in key.lower() or "th" in key.lower():
            if isinstance(value, str):
                if (value in plt.style.available) or (value in scienceplots_styles):
                    plt.style.use(value)
                elif value in sns_themes:
                    sns.set_style(value)
                elif value in sns_contexts:
                    sns.set_context(value)
                else:
                    print(
                        f"\nWarning\n'{value}' is not a plt.style,select on below:\n{plt.style.available+sns_themes+sns_contexts+scienceplots_styles}"
                    )
            if isinstance(value, list):
                for i in value:
                    if (i in plt.style.available) or (i in scienceplots_styles):
                        plt.style.use(i)
                    elif i in sns_themes:
                        sns.set_style(i)
                    elif i in sns_contexts:
                        sns.set_context(i)
                    else:
                        print(
                            f"\nWarning\n'{i}' is not a plt.style,select on below:\n{plt.style.available+sns_themes+sns_contexts+scienceplots_styles}"
                        )
        if "la" in key.lower():
            if "loc" in key.lower() or "po" in key.lower():
                for i in value:
                    if "l" in i.lower():
                        ax.yaxis.set_label_position("left")
                    if "r" in i.lower():
                        ax.yaxis.set_label_position("right")
                    if "t" in i.lower():
                        ax.xaxis.set_label_position("top")
                    if "b" in i.lower():
                        ax.xaxis.set_label_position("bottom")
            if ("x" in key.lower()) and (
                "tic" not in key.lower() and "tk" not in key.lower()
            ):
                ax.set_xlabel(value, fontname=fontname)
            if ("y" in key.lower()) and (
                "tic" not in key.lower() and "tk" not in key.lower()
            ):
                ax.set_ylabel(value, fontname=fontname)
            if ("z" in key.lower()) and (
                "tic" not in key.lower() and "tk" not in key.lower()
            ):
                ax.set_zlabel(value, fontname=fontname)
        # tick location
        if "tic" in key.lower() or "tk" in key.lower():
            if ("loc" in key.lower()) or ("po" in key.lower()):
                if isinstance(value, (str, list)):
                    loc = []
                    for i in value:
                        if ("l" in i.lower()) and ("a" not in i.lower()):
                            ax.yaxis.set_ticks_position("left")
                        if "r" in i.lower():
                            ax.yaxis.set_ticks_position("right")
                        if "t" in i.lower():
                            ax.xaxis.set_ticks_position("top")
                        if "b" in i.lower():
                            ax.xaxis.set_ticks_position("bottom")
                        if i.lower() in ["a", "both", "all", "al", ":"]:
                            ax.xaxis.set_ticks_position("both")
                            ax.yaxis.set_ticks_position("both")
                        if i.lower() in ["xnone",'xoff',"none"]:
                            ax.xaxis.set_ticks_position("none")
                        if i.lower() in ["ynone",'yoff','none']:
                            ax.yaxis.set_ticks_position("none")
            # ticks / labels
            elif "x" in key.lower():
                if "la" not in key.lower():
                    ax.set_xticks(value)
                if "la" in key.lower():
                    ax.set_xticklabels(value)
            elif "y" in key.lower():
                if "la" not in key.lower():
                    ax.set_yticks(value)
                if "la" in key.lower():
                    ax.set_yticklabels(value)
            elif "z" in key.lower():
                if "la" not in key.lower():
                    ax.set_zticks(value)
                if "la" in key.lower():
                    ax.set_zticklabels(value)
        # rotation
        if "angle" in key.lower() or ("rot" in key.lower()):
            if "x" in key.lower():
                ax.tick_params(axis="x", rotation=value)
            if "y" in key.lower():
                ax.tick_params(axis="y", rotation=value)

        if "bo" in key in key:  # and ("p" in key or "l" in key):
                # print("'ticks' style is recommended")
            if isinstance(value, (str, list)):
                locations = []
                for i in value:
                    if "l" in i.lower():
                        locations.append("left")
                    if "r" in i.lower():
                        locations.append("right")
                    if "t" in i.lower():
                        locations.append("top")
                    if "b" in i.lower():
                        locations.append("bottom")
                    if i.lower() in ["a", "both", "all", "al", ":"]:
                        [
                            locations.append(x)
                            for x in ["left", "right", "top", "bottom"]
                        ]
                for i in value:
                    if i.lower() in "none":
                        locations = []
                # check spines
                for loc, spi in ax.spines.items():
                    if loc in locations:
                        spi.set_position(("outward", 0))
                    else:
                        spi.set_color("none")  # no spine
        if key == "tick" or key == "ticks" or key == "ticks_para":
            if isinstance(value, dict):
                for k, val in value.items():
                    if "wh" in k.lower():
                        ax.tick_params(
                            which=val
                        )  # {'major', 'minor', 'both'}, default: 'major'
                    elif "dir" in k.lower():
                        ax.tick_params(direction=val)  # {'in', 'out', 'inout'}
                    elif "len" in k.lower():
                        ax.tick_params(length=val)
                    elif ("wid" in k.lower()) or ("wd" in k.lower()):
                        ax.tick_params(width=val)
                    elif "ax" in k.lower():
                        ax.tick_params(axis=val)  # {'x', 'y', 'both'}, default: 'both'
                    elif ("c" in k.lower()) and ("ect" not in k.lower()):
                        ax.tick_params(colors=val)  # Tick color.
                    elif "pad" in k.lower():
                        ax.tick_params(
                            pad=val
                        )  # float, distance in points between tick and label
                    elif (
                        ("lab" in k.lower())
                        and ("s" in k.lower())
                        and ("z" in k.lower())
                    ):
                        ax.tick_params(
                            labelsize=val
                        )  # float, distance in points between tick and label

        if "mi" in key.lower() and "tic" in key.lower():
            if "x" in value.lower() or "x" in key.lower():
                ax.xaxis.set_minor_locator(tck.AutoMinorLocator())  # ax.minorticks_on()
            if "y" in value.lower() or "y" in key.lower():
                ax.yaxis.set_minor_locator(
                    tck.AutoMinorLocator()
                )  # ax.minorticks_off()
            if value.lower() in ["both", ":", "all", "a", "b", "on"]:
                ax.minorticks_on()
        if key == "colormap" or key == "cmap":
            plt.set_cmap(value)
    def sets_small(ax,key, value):
        if key == "figsize":
            pass
        if key == "xlim":
            ax.set_xlim(value)
        if key == "ylim":
            ax.set_ylim(value)
        if key == "zlim":
            ax.set_zlim(value)
        if "sc" in key.lower():
            if "x" in key.lower():
                ax.set_xscale(value)
            if "y" in key.lower():
                ax.set_yscale(value)
            if "z" in key.lower():
                ax.set_zscale(value)
        if key == "grid":
            if isinstance(value, dict):
                for k, val in value.items():
                    if "wh" in k.lower():
                        ax.grid(
                            which=val
                        )  # {'major', 'minor', 'both'}, default: 'major'
                    elif "ax" in k.lower():
                        ax.grid(axis=val)  # {'x', 'y', 'both'}, default: 'both'
                    elif ("c" in k.lower()) and ("ect" not in k.lower()):
                        ax.grid(color=val)  # Tick color.
                    elif "l" in k.lower() and ("s" in k.lower()):
                        ax.grid(linestyle=val)
                    elif "l" in k.lower() and ("w" in k.lower()):
                        ax.grid(linewidth=val)
                    elif "al" in k.lower():
                        ax.grid(alpha=val)
            else:
                if value == "on" or value is True:
                    ax.grid(visible=True)
                elif value == "off" or value is False:
                    ax.grid(visible=False)
        if "tit" in key.lower():
            if "sup" in key.lower():
                plt.suptitle(value)
            else:
                ax.set_title(value)
        if key.lower() in ["spine", "adjust", "ad", "sp", "spi", "adj","spines"]:
            if isinstance(value, bool) or (value in ["go", "do", "ja", "yes"]):
                if value:
                    adjust_spines(ax)  # dafault distance=2
            if isinstance(value, (float, int)):
                adjust_spines(ax=ax, distance=value)
        if "c" in key.lower() and ("sp" in key.lower() or "ax" in key.lower()):
            for loc, spi in ax.spines.items():
                spi.set_color(value)

    for arg in args:
        if isinstance(arg,matplotlib.axes._axes.Axes):
            ax=arg
            args=args[1:]
    if 'ax' not in locals():
        ax=plt.gca()

    for arg in args:
        if isinstance(arg, dict):
            for k, val in arg.items():
                sets_priority(ax,k, val)
            for k, val in arg.items():
                sets_small(ax,k, val)
        else:
            Nargin = len(args) // 2
            ax.labelFontSizeMultiplier = 1
            ax.titleFontSizeMultiplier = 1
            ax.set_facecolor("w")

            for ip in range(Nargin):
                key = args[ip * 2].lower()
                value = args[ip * 2 + 1]
                sets_priority(ax,key, value)
            for ip in range(Nargin):
                key = args[ip * 2].lower()
                value = args[ip * 2 + 1]
                sets_small(ax,key, value)
    colors = [
        "#474747",
        "#FF2C00",
        "#0C5DA5",
        "#845B97",
        "#58BBCC",
        "#FF9500",
        "#D57DBE",
    ]
    matplotlib.rcParams["axes.prop_cycle"] = cycler(color=colors)
    if len(fig.get_axes()) > 1:
        plt.tight_layout()
        plt.gcf().align_labels()

def thumbnail(dir_img_list,figsize=(10,10),dpi=100, dir_save=None, kind='.png'):
    """
    Display a thumbnail figure of all images in the specified directory.
    Args:
        dir_img_list (list): List of the Directory containing the images.
    """
    num_images = len(dir_img_list)
    if not kind.startswith('.'):
        kind='.'+kind

    if num_images == 0:
        print("No images found to display.")
        return
    grid_size = int(num_images ** 0.5) + 1 # Determine grid size
    fig, axs = plt.subplots(grid_size, grid_size, figsize=figsize,dpi=dpi)
    for ax, image_file in zip(axs.flatten(), dir_img_list):
        try:
            img = Image.open(image_file)
            ax.imshow(img)
            ax.axis('off')
        except:
            continue
    # for ax in axs.flatten():
    #     ax.axis('off')
    [ax.axis("off") for ax in axs.flatten()]
    plt.tight_layout()
    if dir_save is None:
        plt.show()
    else:
        if basename(dir_save):
            fname= basename(dir_save) +kind
        else:
            fname= "_thumbnail_"+basename(dirname(dir_save)[:-1])+'.png'
        if dirname(dir_img_list[0]) == dirname(dir_save):
            figsave(dirname(dir_save[:-1]),fname)
        else:
            figsave(dirname(dir_save),fname)
# usage:
# fpath = "/Users/macjianfeng/Dropbox/github/python/py2ls/tests/xample_netfinder/images/"
# thumbnail(listdir(fpath,'png').fpath.to_list(),dir_save=dirname(fpath))
def read_mplstyle(style_file):
    # Load the style file
    plt.style.use(style_file)

    # Get the current style properties
    style_dict = plt.rcParams

    # Convert to dictionary
    style_dict = dict(style_dict)
    # Print the style dictionary
    for i, j in style_dict.items():
        print(f"\n{i}::::{j}")
    return style_dict
# #example usage:
# style_file = "/ std-colors.mplstyle"
# style_dict = read_mplstyle(style_file)


# search and fine the director of the libary, which installed at local
def dir_lib(lib_oi):
    import site

    # Get the site-packages directory
    f = listdir(site.getsitepackages()[0], "folder")

    # Find Seaborn directory within site-packages
    dir_list = []
    for directory in f.fpath:
        if lib_oi in directory.lower():
            dir_list.append(directory)

    if dir_list != []:
        print(f"{lib_oi} directory:", dir_list)
    else:
        print(f"Cannot find the {lib_oi} in site-packages directory.")
    return dir_list
# example usage:
# dir_lib("seaborn") 

# set up the colorlist, give the number, or the colormap's name
def get_color(n=1, cmap="auto", how="start"):
    # Extract the colormap as a list
    def cmap2hex(cmap_name):
        cmap_ = matplotlib.pyplot.get_cmap(cmap_name)
        colors = [cmap_(i) for i in range(cmap_.N)]
        return [matplotlib.colors.rgb2hex(color) for color in colors]
        # usage: clist = cmap2hex("viridis")
    # cycle times, total number is n (defaultn=10)
    def cycle2list(colorlist, n=10):
        cycler_ = cycler(tmp=colorlist)
        clist = []
        for i, c_ in zip(range(n), cycler_()):
            clist.append(c_["tmp"])
            if i > n:
                break
        return clist
    def hue2rgb(hex_colors):
        def hex_to_rgb(hex_color):
            """Converts a hexadecimal color code to RGB values."""
            if hex_colors.startswith("#"):
                hex_color = hex_color.lstrip("#")
            return tuple(int(hex_color[i : i + 2], 16) / 255.0 for i in (0, 2, 4))
        if isinstance(hex_colors, str):
            return hex_to_rgb(hex_colors)
        elif isinstance(hex_colors, (list)):
            """Converts a list of hexadecimal color codes to a list of RGB values."""
            rgb_values = [hex_to_rgb(hex_color) for hex_color in hex_colors]
            return rgb_values
    if "aut" in cmap:
        colorlist = [
            "#474747",
            "#FF2C00",
            "#0C5DA5",
            "#845B97",
            "#58BBCC",
            "#FF9500",
            "#D57DBE",
        ]
    else:
        colorlist = cmap2hex(cmap)
    if "st" in how.lower() or "be" in how.lower():
        # cycle it
        clist = cycle2list(colorlist, n=n)
    if "l" in how.lower() or "p" in how.lower():
        clist = []
        [
            clist.append(colorlist[i])
            for i in [int(i) for i in np.linspace(0, len(colorlist) - 1, n)]
        ]

    return clist  # a color list
    # example usage: clist = get_color(4,cmap="auto", how="start") # get_color(4, cmap="hot", how="linspace")

""" 
    # n = 7
    # clist = get_color(n, cmap="auto", how="linspace")  # get_color(100)
    # plt.figure(figsize=[8, 5], dpi=100)
    # x = np.linspace(0, 2 * np.pi, 50) * 100
    # y = np.sin(x)
    # for i in range(1, n + 1):
    #     plt.plot(x, y + i, c=clist[i - 1], lw=5, label=str(i))
    # plt.legend()
    # plt.ylim(-2, 20)
    # figsets(plt.gca(), {"style": "whitegrid"}) """


class FileInfo:
    def __init__(self, size, creation_time, ctime, mod_time, mtime, parent_dir, fname, kind, extra_info=None):
        self.size = size
        self.creation_time = creation_time
        self.ctime = ctime
        self.mod_time = mod_time
        self.mtime = mtime
        self.parent_dir = parent_dir
        self.fname = fname
        self.kind = kind
        if extra_info:
            for key, value in extra_info.items():
                setattr(self, key, value)
        print("to show the res: 'finfo(fpath).show()'")

    def __repr__(self):
        return (f"FileInfo(size={self.size} MB, creation_time='{self.creation_time}', "
                f"ctime='{self.ctime}', mod_time='{self.mod_time}', mtime='{self.mtime}', "
                f"parent_dir='{self.parent_dir}', fname='{self.fname}', kind='{self.kind}')")

    def __str__(self):
        return (f"FileInfo:\n"
                f"  Size: {self.size} MB\n"
                f"  Creation Time: {self.creation_time}\n"
                f"  CTime: {self.ctime}\n"
                f"  Modification Time: {self.mod_time}\n"
                f"  MTime: {self.mtime}\n"
                f"  Parent Directory: {self.parent_dir}\n"
                f"  File Name: {self.fname}\n"
                f"  Kind: {self.kind}")
    def show(self):
        # Convert the object to a dictionary
        return {
            "size": self.size,
            "creation_time": self.creation_time,
            "ctime": self.ctime,
            "mod_time": self.mod_time,
            "mtime": self.mtime,
            "parent_dir": self.parent_dir,
            "fname": self.fname,
            "kind": self.kind,
            **{key: getattr(self, key) for key in vars(self) if key not in ["size", "creation_time", "ctime", "mod_time", "mtime", "parent_dir", "fname", "kind"]}
        }

def finfo(fpath):
    fname, fmt = os.path.splitext(fpath)
    dir_par = os.path.dirname(fpath) + '/'
    data = {
        "size": round(os.path.getsize(fpath) / 1024 / 1024, 3),
        "creation_time": time.ctime(os.path.getctime(fpath)),
        "ctime": time.ctime(os.path.getctime(fpath)),
        "mod_time": time.ctime(os.path.getmtime(fpath)),
        "mtime": time.ctime(os.path.getmtime(fpath)),
        "parent_dir": dir_par,
        "fname": fname.replace(dir_par, ""),
        "kind": fmt
    }
    extra_info = {}
    if data["kind"] == ".pdf":
        extra_info = pdfinfo_from_path(fpath)
    
    return FileInfo(
        size=data["size"],
        creation_time=data["creation_time"],
        ctime=data["ctime"],
        mod_time=data["mod_time"],
        mtime=data["mtime"],
        parent_dir=data["parent_dir"],
        fname=data["fname"],
        kind=data["kind"],
        extra_info=extra_info
    )